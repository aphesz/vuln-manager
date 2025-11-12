"""
Generate default DOCX module templates programmatically.

Run: python -m app.report_modules.generate_templates
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MODULE_DIR = Path(__file__).parent


def create_title_page_template():
    """Create title page module with company branding placeholders."""
    doc = Document()
    
    # Spacer
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = doc.add_heading("Security Assessment Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project name
    proj = doc.add_paragraph("{{ project.name }}")
    proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj.runs[0].font.size = Pt(18)
    proj.runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    
    meta_table.cell(0, 0).text = "Report Date:"
    meta_table.cell(0, 1).text = "{{ report_date }}"
    
    meta_table.cell(1, 0).text = "Consultant:"
    meta_table.cell(1, 1).text = "{{ project.consultant_name }}"
    
    meta_table.cell(2, 0).text = "Company:"
    meta_table.cell(2, 1).text = "{{ company_name }}"
    
    meta_table.cell(3, 0).text = "Total Findings:"
    meta_table.cell(3, 1).text = "{{ total_findings }}"
    
    meta_table.cell(4, 0).text = "Assessment Period:"
    meta_table.cell(4, 1).text = "{{ assessment_period }}"
    
    # Page break
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "title_page.docx")
    print("✓ Created title_page.docx")


def create_executive_summary_template():
    """Create executive summary module."""
    doc = Document()
    
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph()
    
    # Summary paragraph
    summary = doc.add_paragraph(
        "This security assessment of "
    )
    summary.add_run("{{ project.name }}").bold = True
    summary.add_run(" identified ")
    summary.add_run("{{ total_findings }}").bold = True
    summary.add_run(" security findings across various risk levels. ")
    summary.add_run("{{ critical_count }}").bold = True
    summary.add_run(" Critical and ")
    summary.add_run("{{ high_count }}").bold = True
    summary.add_run(" High severity findings require immediate attention.")
    
    doc.add_paragraph()
    
    # Risk summary heading
    doc.add_heading("Risk Summary", level=2)
    
    # Risk counts table
    risk_table = doc.add_table(rows=6, cols=2)
    risk_table.style = 'Light Grid Accent 1'
    
    risk_table.cell(0, 0).text = "Risk Level"
    risk_table.cell(0, 1).text = "Count"
    risk_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    risk_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
    risk_table.cell(1, 0).text = "Critical"
    risk_table.cell(1, 1).text = "{{ critical_count }}"
    
    risk_table.cell(2, 0).text = "High"
    risk_table.cell(2, 1).text = "{{ high_count }}"
    
    risk_table.cell(3, 0).text = "Medium"
    risk_table.cell(3, 1).text = "{{ medium_count }}"
    
    risk_table.cell(4, 0).text = "Low"
    risk_table.cell(4, 1).text = "{{ low_count }}"
    
    risk_table.cell(5, 0).text = "Informational"
    risk_table.cell(5, 1).text = "{{ informational_count }}"
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "executive_summary.docx")
    print("✓ Created executive_summary.docx")


def create_detailed_findings_template():
    """Create detailed findings module with full field support."""
    doc = Document()
    
    doc.add_heading("Detailed Findings", level=1)
    doc.add_paragraph()
    
    # Findings loop
    doc.add_paragraph("{% for f in findings %}")
    
    # Finding heading
    doc.add_heading("{{ f.section_number }} {{ f.title }} ({{ f.risk_rating }})", level=2)
    
    # Create finding table
    finding_table = doc.add_table(rows=15, cols=2)
    finding_table.style = 'Light Grid Accent 1'
    finding_table.columns[0].width = Inches(2.0)
    finding_table.columns[1].width = Inches(4.5)
    
    # Row 0: Risk Rating
    finding_table.cell(0, 0).text = "Risk Rating:"
    finding_table.cell(0, 1).text = "{{ f.risk_rating }}"
    finding_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    
    # Row 1: Instances
    finding_table.cell(1, 0).text = "Instances:"
    finding_table.cell(1, 1).text = "{{ f.instances_count }}"
    
    # Row 2: Status
    finding_table.cell(2, 0).text = "Issue Status:"
    finding_table.cell(2, 1).text = "{{ f.issue_status }}"
    
    # Row 3: Review Status
    finding_table.cell(3, 0).text = "Review Status:"
    finding_table.cell(3, 1).text = "{{ f.review_status }} ({{ f.reviewer_name }})"
    
    # Row 4: SLA
    finding_table.cell(4, 0).text = "SLA Status:"
    finding_table.cell(4, 1).text = "{{ f.sla_status }} - Due: {{ f.remediation_deadline }}"
    
    # Row 5: Remediation Owner
    finding_table.cell(5, 0).text = "Remediation Owner:"
    finding_table.cell(5, 1).text = "{{ f.remediation_owner }}"
    
    # Row 6: Jira
    finding_table.cell(6, 0).text = "Jira Ticket:"
    finding_table.cell(6, 1).text = "{{ f.jira_issue_key }} ({{ f.jira_status }})"
    
    # Row 7: CVE/CWE
    finding_table.cell(7, 0).text = "CVE / CWE:"
    finding_table.cell(7, 1).text = "{{ f.cve_id }} / {{ f.cwe_id }}"
    
    # Row 8: CVSS
    finding_table.cell(8, 0).text = "CVSS Score:"
    finding_table.cell(8, 1).text = "{{ f.cvss_score }} ({{ f.cvss_vector }})"
    
    # Row 9: OWASP
    finding_table.cell(9, 0).text = "OWASP Category:"
    finding_table.cell(9, 1).text = "{{ f.owasp_category }} - Risk: {{ f.owasp_risk_rating }}"
    
    # Row 10: Affected Resources
    finding_table.cell(10, 0).text = "Affected Resources:"
    finding_table.cell(10, 1).text = "{{ f.affected_resources }}"
    
    # Row 11: Timeline
    finding_table.cell(11, 0).text = "Timeline:"
    finding_table.cell(11, 1).text = "Discovered: {{ f.discovered_at }} | Resolved: {{ f.resolved_at }}"
    
    # Row 12: Description header
    finding_table.cell(12, 0).text = "Description:"
    finding_table.cell(12, 0).paragraphs[0].runs[0].font.bold = True
    finding_table.cell(12, 1).text = ""
    
    # Row 13: Description content
    finding_table.cell(13, 0).text = ""
    finding_table.cell(13, 1).text = "{{ f.description_text }}"
    
    # Row 14: Impact
    finding_table.cell(14, 0).text = "Impact:"
    finding_table.cell(14, 0).paragraphs[0].runs[0].font.bold = True
    finding_table.cell(14, 1).text = "{{ f.impact }}"
    
    doc.add_paragraph()
    
    # Remediation section
    doc.add_heading("Remediation", level=3)
    doc.add_paragraph("{{ f.remediation_text }}")
    
    # POC section
    doc.add_heading("Proof of Concept", level=3)
    doc.add_paragraph("{{ f.poc_content }}")
    
    # References
    doc.add_paragraph()
    ref = doc.add_paragraph("References: ")
    ref.add_run("{{ f.references_url }}").font.color.rgb = RGBColor(0, 0, 255)
    
    doc.add_paragraph()
    doc.add_paragraph("{% endfor %}")
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "detailed_findings.docx")
    print("✓ Created detailed_findings.docx")


def create_recommendations_template():
    """Create recommendations module."""
    doc = Document()
    
    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph()
    
    # Critical priority
    doc.add_heading("Critical Priority", level=2)
    doc.add_paragraph(
        "{% if critical_count > 0 %}"
        "Address all {{ critical_count }} Critical severity findings immediately. "
        "These vulnerabilities pose a direct threat to system security and should be "
        "remediated within 24-48 hours."
        "{% else %}"
        "No critical findings identified."
        "{% endif %}"
    )
    
    # High priority
    doc.add_heading("High Priority", level=2)
    doc.add_paragraph(
        "{% if high_count > 0 %}"
        "Remediate {{ high_count }} High severity findings within 1-2 weeks. "
        "These issues present significant security risks and should be prioritized "
        "in the remediation plan."
        "{% else %}"
        "No high severity findings identified."
        "{% endif %}"
    )
    
    # General recommendations
    doc.add_heading("General Security Recommendations", level=2)
    
    doc.add_paragraph("Conduct regular security assessments (quarterly recommended) to identify "
                     "new vulnerabilities and validate remediation efforts.", style='List Bullet')
    
    doc.add_paragraph("Provide security awareness training to development teams to prevent "
                     "common vulnerabilities from being introduced during the development lifecycle.",
                     style='List Bullet')
    
    doc.add_paragraph("Implement automated security scanning in CI/CD pipelines for continuous "
                     "vulnerability detection.", style='List Bullet')
    
    doc.add_paragraph("Establish a formal vulnerability disclosure and remediation process with "
                     "defined SLAs by severity level.", style='List Bullet')
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "recommendations.docx")
    print("✓ Created recommendations.docx")


def create_top_findings_template():
    """Create top findings summary module."""
    doc = Document()
    
    doc.add_heading("Top Priority Findings", level=1)
    doc.add_paragraph()
    
    intro = doc.add_paragraph(
        "The following table summarizes the highest priority findings that require "
        "immediate attention."
    )
    
    doc.add_paragraph()
    
    # Top findings table
    top_table = doc.add_table(rows=1, cols=5)
    top_table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = top_table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Finding"
    hdr_cells[2].text = "Risk"
    hdr_cells[3].text = "Instances"
    hdr_cells[4].text = "SLA Status"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows (Jinja loop) - using tr template
    doc.add_paragraph("{%tr for f in findings[:10] %}")
    
    row_template = top_table.add_row().cells
    row_template[0].text = "{{ loop.index }}"
    row_template[1].text = "{{ f.title[:60] }}"
    row_template[2].text = "{{ f.risk_rating }}"
    row_template[3].text = "{{ f.instances_count }}"
    row_template[4].text = "{{ f.sla_status or 'N/A' }}"
    
    doc.add_paragraph("{%endtr%}")
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "top_findings.docx")
    print("✓ Created top_findings.docx")


def create_sla_status_template():
    """Create SLA tracking module."""
    doc = Document()
    
    doc.add_heading("SLA Tracking Report", level=1)
    doc.add_paragraph()
    
    # Summary
    summary = doc.add_paragraph(
        "This report tracks the remediation SLA status for all findings. "
    )
    summary.add_run("{{ overdue_count }}").bold = True
    summary.add_run(" findings are currently overdue, and ")
    summary.add_run("{{ at_risk_count }}").bold = True
    summary.add_run(" are at risk of missing their SLA deadlines.")
    
    doc.add_paragraph()
    
    # SLA status table
    doc.add_heading("Findings by SLA Status", level=2)
    
    sla_table = doc.add_table(rows=1, cols=6)
    sla_table.style = 'Light Grid Accent 1'
    
    hdr = sla_table.rows[0].cells
    hdr[0].text = "Finding"
    hdr[1].text = "Risk"
    hdr[2].text = "SLA Status"
    hdr[3].text = "Deadline"
    hdr[4].text = "Owner"
    hdr[5].text = "Jira"
    
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph("{% for f in findings %}")
    
    row = sla_table.add_row().cells
    row[0].text = "{{ f.title[:40] }}"
    row[1].text = "{{ f.risk_rating }}"
    row[2].text = "{{ f.sla_status }}"
    row[3].text = "{{ f.remediation_deadline }}"
    row[4].text = "{{ f.remediation_owner }}"
    row[5].text = "{{ f.jira_issue_key }}"
    
    doc.add_paragraph("{% endfor %}")
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "sla_status.docx")
    print("✓ Created sla_status.docx")


def create_risk_charts_template():
    """Create risk distribution charts module."""
    doc = Document()
    
    doc.add_heading("Risk Distribution Analysis", level=1)
    doc.add_paragraph()
    
    # Introduction
    intro = doc.add_paragraph(
        "The following charts visualize the distribution of security findings "
        "across different risk categories, compliance frameworks, and timelines."
    )
    
    doc.add_paragraph()
    
    # Risk by severity chart placeholder
    doc.add_heading("Findings by Risk Level", level=2)
    chart_note = doc.add_paragraph(
        "[CHART: Risk Distribution - Critical: {{ critical_count }}, "
        "High: {{ high_count }}, Medium: {{ medium_count }}, Low: {{ low_count }}]"
    )
    chart_note.runs[0].italic = True
    
    doc.add_paragraph()
    
    # Risk breakdown table
    risk_table = doc.add_table(rows=1, cols=3)
    risk_table.style = 'Light Grid Accent 1'
    
    hdr_cells = risk_table.rows[0].cells
    hdr_cells[0].text = "Risk Level"
    hdr_cells[1].text = "Count"
    hdr_cells[2].text = "Percentage"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for level in ["Critical", "High", "Medium", "Low", "Informational"]:
        row = risk_table.add_row().cells
        row[0].text = f"{level}"
        row[1].text = "{{ " + level.lower() + "_count }}"
        row[2].text = "{{ " + level.lower() + "_percentage }}%"
    
    doc.add_paragraph()
    
    # OWASP distribution
    doc.add_heading("OWASP Top 10 Distribution", level=2)
    owasp_note = doc.add_paragraph(
        "[CHART: OWASP Category Distribution]"
    )
    owasp_note.runs[0].italic = True
    
    doc.add_paragraph()
    
    # CWE distribution
    doc.add_heading("Top CWE Categories", level=2)
    cwe_note = doc.add_paragraph(
        "[CHART: Top 10 CWE Classifications]"
    )
    cwe_note.runs[0].italic = True
    
    doc.add_paragraph()
    
    # Timeline chart
    doc.add_heading("Finding Discovery Timeline", level=2)
    timeline_note = doc.add_paragraph(
        "[CHART: Findings discovered over time]"
    )
    timeline_note.runs[0].italic = True
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "risk_charts.docx")
    print("✓ Created risk_charts.docx")


def create_appendix_template():
    """Create technical appendix module."""
    doc = Document()
    
    doc.add_heading("Appendix", level=1)
    doc.add_paragraph()
    
    # A. Methodology
    doc.add_heading("A. Testing Methodology", level=2)
    doc.add_paragraph(
        "This security assessment was conducted using a combination of automated "
        "scanning tools and manual testing techniques. The following scanners and "
        "methodologies were employed:"
    )
    
    doc.add_paragraph("Burp Suite Professional - Web application vulnerability scanning", style='List Bullet')
    doc.add_paragraph("Nessus Professional - Network and infrastructure scanning", style='List Bullet')
    doc.add_paragraph("Manual code review and penetration testing", style='List Bullet')
    doc.add_paragraph("OWASP Testing Guide methodologies", style='List Bullet')
    
    doc.add_paragraph()
    
    # B. Scope
    doc.add_heading("B. Assessment Scope", level=2)
    
    scope_table = doc.add_table(rows=1, cols=2)
    scope_table.style = 'Light Grid Accent 1'
    
    hdr_cells = scope_table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Details"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Scope details
    scope_items = [
        ("Project Name", "{{ project_name }}"),
        ("Assessment Period", "{{ assessment_period }}"),
        ("Total Findings", "{{ total_findings }}"),
        ("Unique Vulnerabilities", "{{ findings_count }}"),
        ("Affected Components", "{{ affected_count }}"),
        ("Consultant", "{{ consultant_name }}")
    ]
    
    for item, value in scope_items:
        row = scope_table.add_row().cells
        row[0].text = item
        row[1].text = value
    
    doc.add_paragraph()
    
    # C. Risk Rating Criteria
    doc.add_heading("C. Risk Rating Criteria", level=2)
    doc.add_paragraph(
        "Findings are classified using the following risk rating criteria:"
    )
    
    risk_criteria = doc.add_table(rows=1, cols=3)
    risk_criteria.style = 'Light Grid Accent 1'
    
    hdr = risk_criteria.rows[0].cells
    hdr[0].text = "Rating"
    hdr[1].text = "CVSS Score"
    hdr[2].text = "Remediation SLA"
    
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
    
    criteria = [
        ("Critical", "9.0 - 10.0", "24-48 hours"),
        ("High", "7.0 - 8.9", "1-2 weeks"),
        ("Medium", "4.0 - 6.9", "1 month"),
        ("Low", "0.1 - 3.9", "3 months"),
        ("Informational", "0.0", "Best effort")
    ]
    
    for rating, cvss, sla in criteria:
        row = risk_criteria.add_row().cells
        row[0].text = rating
        row[1].text = cvss
        row[2].text = sla
    
    doc.add_paragraph()
    
    # D. References
    doc.add_heading("D. References and Resources", level=2)
    doc.add_paragraph("OWASP Top 10: https://owasp.org/www-project-top-ten/", style='List Bullet')
    doc.add_paragraph("CWE Top 25: https://cwe.mitre.org/top25/", style='List Bullet')
    doc.add_paragraph("CVSS Calculator: https://www.first.org/cvss/calculator/", style='List Bullet')
    doc.add_paragraph("NIST Vulnerability Database: https://nvd.nist.gov/", style='List Bullet')
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "appendix.docx")
    print("✓ Created appendix.docx")


def create_compliance_owasp_template():
    """Create OWASP Top 10 compliance mapping module."""
    doc = Document()
    
    doc.add_heading("OWASP Top 10 Compliance Report", level=1)
    doc.add_paragraph()
    
    intro = doc.add_paragraph(
        "This section maps identified vulnerabilities to the OWASP Top 10 2021 "
        "framework, providing visibility into application security posture across "
        "industry-standard risk categories."
    )
    
    doc.add_paragraph()
    
    # Summary table
    doc.add_heading("OWASP Top 10 Coverage", level=2)
    
    owasp_table = doc.add_table(rows=1, cols=4)
    owasp_table.style = 'Light Grid Accent 1'
    
    hdr_cells = owasp_table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Findings"
    hdr_cells[3].text = "Highest Risk"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # OWASP Top 10 categories
    owasp_categories = [
        ("A01:2021", "Broken Access Control", "{{ a01_count }}", "{{ a01_max_risk }}"),
        ("A02:2021", "Cryptographic Failures", "{{ a02_count }}", "{{ a02_max_risk }}"),
        ("A03:2021", "Injection", "{{ a03_count }}", "{{ a03_max_risk }}"),
        ("A04:2021", "Insecure Design", "{{ a04_count }}", "{{ a04_max_risk }}"),
        ("A05:2021", "Security Misconfiguration", "{{ a05_count }}", "{{ a05_max_risk }}"),
        ("A06:2021", "Vulnerable Components", "{{ a06_count }}", "{{ a06_max_risk }}"),
        ("A07:2021", "ID & Auth Failures", "{{ a07_count }}", "{{ a07_max_risk }}"),
        ("A08:2021", "Software & Data Integrity", "{{ a08_count }}", "{{ a08_max_risk }}"),
        ("A09:2021", "Logging Failures", "{{ a09_count }}", "{{ a09_max_risk }}"),
        ("A10:2021", "SSRF", "{{ a10_count }}", "{{ a10_max_risk }}")
    ]
    
    for cat_id, desc, count, risk in owasp_categories:
        row = owasp_table.add_row().cells
        row[0].text = cat_id
        row[1].text = desc
        row[2].text = count
        row[3].text = risk
    
    doc.add_paragraph()
    
    # Key findings
    doc.add_heading("Key OWASP Findings", level=2)
    doc.add_paragraph(
        "The most critical OWASP categories identified in this assessment:"
    )
    
    doc.add_paragraph()
    doc.add_paragraph("{{ owasp_summary }}")
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "compliance_owasp.docx")
    print("✓ Created compliance_owasp.docx")


def create_compliance_cwe_template():
    """Create CWE classification compliance module."""
    doc = Document()
    
    doc.add_heading("CWE Classification Report", level=1)
    doc.add_paragraph()
    
    intro = doc.add_paragraph(
        "This section classifies identified vulnerabilities according to the "
        "Common Weakness Enumeration (CWE) framework maintained by MITRE. "
        "CWE provides a standardized taxonomy for software weaknesses."
    )
    
    doc.add_paragraph()
    
    # Top CWEs table
    doc.add_heading("Top CWE Classifications", level=2)
    
    cwe_table = doc.add_table(rows=1, cols=4)
    cwe_table.style = 'Light Grid Accent 1'
    
    hdr_cells = cwe_table.rows[0].cells
    hdr_cells[0].text = "CWE ID"
    hdr_cells[1].text = "Weakness Name"
    hdr_cells[2].text = "Count"
    hdr_cells[3].text = "Highest Risk"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Common CWEs (top 10 placeholders)
    for i in range(1, 11):
        row = cwe_table.add_row().cells
        row[0].text = f"{{{{ cwe_{i}_id }}}}"
        row[1].text = f"{{{{ cwe_{i}_name }}}}"
        row[2].text = f"{{{{ cwe_{i}_count }}}}"
        row[3].text = f"{{{{ cwe_{i}_max_risk }}}}"
    
    doc.add_paragraph()
    
    # CWE Top 25 alignment
    doc.add_heading("CWE Top 25 Alignment", level=2)
    doc.add_paragraph(
        "Findings aligned with MITRE's CWE Top 25 Most Dangerous Software Weaknesses:"
    )
    
    doc.add_paragraph()
    doc.add_paragraph("Total CWE Top 25 findings: {{ cwe_top25_count }}")
    doc.add_paragraph("Critical severity: {{ cwe_top25_critical }}")
    doc.add_paragraph("High severity: {{ cwe_top25_high }}")
    
    doc.add_paragraph()
    
    # CVE mapping
    doc.add_heading("CVE References", level=2)
    doc.add_paragraph(
        "Findings with associated Common Vulnerabilities and Exposures (CVE) identifiers:"
    )
    
    doc.add_paragraph()
    doc.add_paragraph("Total CVE-mapped findings: {{ cve_count }}")
    doc.add_paragraph("Latest CVE year: {{ latest_cve_year }}")
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "compliance_cwe.docx")
    print("✓ Created compliance_cwe.docx")


def create_jira_integration_template():
    """Create Jira ticket integration summary module."""
    doc = Document()
    
    doc.add_heading("Jira Integration Report", level=1)
    doc.add_paragraph()
    
    intro = doc.add_paragraph(
        "This section summarizes findings that have been exported to Jira for "
        "tracking and remediation. It provides visibility into ticket status, "
        "assignments, and progress."
    )
    
    doc.add_paragraph()
    
    # Summary stats
    doc.add_heading("Jira Ticket Summary", level=2)
    
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Count"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    metrics = [
        ("Total Findings", "{{ total_findings }}"),
        ("Findings with Jira Tickets", "{{ jira_linked_count }}"),
        ("Open Tickets", "{{ jira_open_count }}"),
        ("In Progress", "{{ jira_in_progress_count }}"),
        ("Resolved", "{{ jira_resolved_count }}"),
        ("Closed", "{{ jira_closed_count }}")
    ]
    
    for metric, value in metrics:
        row = summary_table.add_row().cells
        row[0].text = metric
        row[1].text = value
    
    doc.add_paragraph()
    
    # Ticket details table
    doc.add_heading("Jira Ticket Details", level=2)
    
    ticket_table = doc.add_table(rows=1, cols=5)
    ticket_table.style = 'Light Grid Accent 1'
    
    hdr_cells = ticket_table.rows[0].cells
    hdr_cells[0].text = "Finding"
    hdr_cells[1].text = "Jira Key"
    hdr_cells[2].text = "Status"
    hdr_cells[3].text = "Priority"
    hdr_cells[4].text = "Assignee"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph("{% for f in jira_findings %}")
    
    row_template = ticket_table.add_row().cells
    row_template[0].text = "{{ f.title[:40] }}"
    row_template[1].text = "{{ f.jira_issue_key }}"
    row_template[2].text = "{{ f.jira_status }}"
    row_template[3].text = "{{ f.risk_rating }}"
    row_template[4].text = "{{ f.remediation_owner }}"
    
    doc.add_paragraph("{% endfor %}")
    
    doc.add_paragraph()
    
    # Notes
    doc.add_heading("Integration Notes", level=2)
    doc.add_paragraph(
        "Jira tickets are automatically synchronized with finding status. "
        "Updates to ticket status in Jira are reflected in the vulnerability "
        "management system."
    )
    
    doc.add_paragraph()
    doc.add_paragraph("Last synchronization: {{ last_jira_sync }}")
    doc.add_paragraph("Jira project: {{ jira_project_key }}")
    
    doc.add_page_break()
    
    doc.save(MODULE_DIR / "jira_integration.docx")
    print("✓ Created jira_integration.docx")


def main():
    """Generate all default module templates."""
    print("Generating default report module templates...")
    print()
    
    create_title_page_template()
    create_executive_summary_template()
    create_detailed_findings_template()
    create_recommendations_template()
    create_top_findings_template()
    create_sla_status_template()
    create_risk_charts_template()
    create_appendix_template()
    create_compliance_owasp_template()
    create_compliance_cwe_template()
    create_jira_integration_template()
    
    print()
    print(f"✅ Successfully generated 11 module templates in {MODULE_DIR}")
    print()
    print("Next steps:")
    print("1. Review the generated templates in Word/LibreOffice")
    print("2. Customize layouts, fonts, and styles as needed")
    print("3. Use the /projects/{id}/report/assemble endpoint to generate modular reports")


if __name__ == "__main__":
    main()
