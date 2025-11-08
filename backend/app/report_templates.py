# backend/app/report_templates.py
"""
Advanced Report Template Engine
Generates professional reports in multiple formats with customization support.
"""

from typing import List, Dict, Any, Optional
from datetime import datetime
from sqlmodel import Session, select
from app.models import (
    Project, Finding, Instance, ReportTemplateType, 
    ReportFormat, ReportBranding, FindingBase
)
from app.executive import ExecutiveMetrics
from app.timezone_utils import get_utc_now
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import html
import re
import base64


class ReportTemplateEngine:
    """Engine for generating reports from templates with customization."""
    
    def __init__(self, session: Session, branding: Optional[ReportBranding] = None):
        self.session = session
        self.branding = branding or self._get_default_branding()
    
    def _get_default_branding(self) -> ReportBranding:
        """Get branding settings or return defaults."""
        branding = self.session.exec(select(ReportBranding)).first()
        if not branding:
            # Create default branding
            branding = ReportBranding(
                company_name="VulnManager",
                primary_color="#1976d2",
                secondary_color="#dc004e",
                footer_text="Confidential - Security Assessment Report",
                created_at=get_utc_now(),
                updated_at=get_utc_now()
            )
        return branding
    
    def generate_report(
        self,
        template_type: ReportTemplateType,
        format: ReportFormat,
        project_ids: List[int] = None,
        start_date: Optional[datetime] = None,
        end_date: Optional[datetime] = None,
        include_sections: Optional[List[str]] = None
    ) -> str:
        """
        Generate a report based on template type and format.
        
        Returns:
            File path to generated report
        """
        # Get projects
        if project_ids:
            projects = self.session.exec(
                select(Project).where(Project.id.in_(project_ids))
            ).all()
        else:
            projects = self.session.exec(
                select(Project).where(Project.is_archived == False)
            ).all()
        
        # Generate based on template type
        if template_type == ReportTemplateType.ExecutiveSummary:
            return self._generate_executive_summary(format, projects)
        elif template_type == ReportTemplateType.TechnicalFindings:
            return self._generate_technical_findings(format, projects)
        elif template_type == ReportTemplateType.RiskAssessment:
            return self._generate_risk_assessment(format, projects)
        elif template_type == ReportTemplateType.RemediationStatus:
            return self._generate_remediation_status(format, projects)
        elif template_type == ReportTemplateType.PortfolioOverview:
            return self._generate_portfolio_overview(format, projects)
        else:
            raise ValueError(f"Unsupported template type: {template_type}")
    
    def _generate_executive_summary(
        self, 
        format: ReportFormat, 
        projects: List[Project]
    ) -> str:
        """Generate Executive Summary report."""
        # Get executive metrics
        summary = ExecutiveMetrics.get_executive_summary(self.session)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == ReportFormat.HTML:
            file_path = f"/tmp/executive_summary_{timestamp}.html"
            self._generate_executive_summary_html(summary, projects, file_path)
        elif format == ReportFormat.DOCX:
            file_path = f"/tmp/executive_summary_{timestamp}.docx"
            self._generate_executive_summary_docx(summary, projects, file_path)
        elif format == ReportFormat.PDF:
            file_path = f"/tmp/executive_summary_{timestamp}.pdf"
            self._generate_executive_summary_pdf(summary, projects, file_path)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        return file_path
    
    def _generate_executive_summary_html(
        self, 
        summary: Dict[str, Any], 
        projects: List[Project],
        file_path: str
    ):
        """Generate interactive HTML Executive Summary."""
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Executive Summary - Security Report</title>
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f5f5f5;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: white;
        }}
        header {{
            background: {self.branding.primary_color};
            color: white;
            padding: 30px;
            text-align: center;
            margin-bottom: 30px;
        }}
        h1 {{
            font-size: 2.5em;
            margin-bottom: 10px;
        }}
        .subtitle {{
            font-size: 1.2em;
            opacity: 0.9;
        }}
        .meta {{
            background: #f9f9f9;
            padding: 15px;
            border-left: 4px solid {self.branding.primary_color};
            margin-bottom: 30px;
        }}
        .kpi-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-bottom: 40px;
        }}
        .kpi-card {{
            background: white;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            padding: 20px;
            text-align: center;
            transition: transform 0.2s;
        }}
        .kpi-card:hover {{
            transform: translateY(-5px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }}
        .kpi-value {{
            font-size: 3em;
            font-weight: bold;
            color: {self.branding.primary_color};
            margin: 10px 0;
        }}
        .kpi-label {{
            color: #666;
            font-size: 0.9em;
            text-transform: uppercase;
        }}
        .alert {{
            background: #fff3cd;
            border-left: 4px solid #ffc107;
            padding: 15px;
            margin-bottom: 20px;
        }}
        .alert.critical {{
            background: #f8d7da;
            border-color: #dc3545;
        }}
        .section {{
            margin-bottom: 40px;
        }}
        .section-title {{
            font-size: 1.8em;
            color: {self.branding.primary_color};
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid #e0e0e0;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }}
        th {{
            background: {self.branding.primary_color};
            color: white;
            font-weight: 600;
        }}
        tr:hover {{
            background: #f5f5f5;
        }}
        .badge {{
            display: inline-block;
            padding: 4px 12px;
            border-radius: 12px;
            font-size: 0.85em;
            font-weight: 600;
            color: white;
        }}
        .badge.critical {{ background: #d32f2f; }}
        .badge.high {{ background: #f57c00; }}
        .badge.medium {{ background: #fbc02d; color: #333; }}
        .badge.low {{ background: #388e3c; }}
        .badge.green {{ background: #4caf50; }}
        .badge.yellow {{ background: #ffc107; color: #333; }}
        .badge.orange {{ background: #ff9800; }}
        .badge.red {{ background: #f44336; }}
        .chart-container {{
            position: relative;
            height: 300px;
            margin: 20px 0;
        }}
        footer {{
            text-align: center;
            padding: 20px;
            color: #666;
            border-top: 1px solid #ddd;
            margin-top: 40px;
        }}
        @media print {{
            body {{ background: white; }}
            .container {{ box-shadow: none; }}
            .kpi-card:hover {{ transform: none; box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>🛡️ Executive Summary</h1>
            <div class="subtitle">Security Posture & Risk Assessment</div>
        </header>
        
        <div class="meta">
            <strong>Report Generated:</strong> {summary['generated_at']}<br>
            <strong>Company:</strong> {self.branding.company_name or 'VulnManager'}<br>
            <strong>Coverage:</strong> {summary['total_projects']} Active Projects
        </div>
        
        {self._generate_alert_html(summary)}
        
        <div class="section">
            <h2 class="section-title">📊 Key Performance Indicators</h2>
            <div class="kpi-grid">
                <div class="kpi-card">
                    <div class="kpi-label">Active Projects</div>
                    <div class="kpi-value">{summary['total_projects']}</div>
                </div>
                <div class="kpi-card">
                    <div class="kpi-label">Total Findings</div>
                    <div class="kpi-value">{summary['total_findings']}</div>
                </div>
                <div class="kpi-card">
                    <div class="kpi-label">MTTR (Days)</div>
                    <div class="kpi-value">{summary['mttr_days']}</div>
                </div>
                <div class="kpi-card">
                    <div class="kpi-label">Trend</div>
                    <div class="kpi-value" style="font-size: 2em;">{self._get_trend_emoji(summary['trend']['trend_direction'])}</div>
                    <div style="font-size: 0.9em; color: #666;">{summary['trend']['trend_direction'].title()}</div>
                </div>
            </div>
        </div>
        
        {self._generate_findings_chart_html(summary)}
        {self._generate_compliance_html(summary)}
        {self._generate_top_risks_table_html(summary)}
        
        <footer>
            <p>{self.branding.footer_text}</p>
            <p style="font-size: 0.85em; margin-top: 10px;">Generated by VulnManager - Advanced Security Assessment Platform</p>
        </footer>
    </div>
    
    <script>
        // Findings by Severity Chart
        const ctx = document.getElementById('findingsChart').getContext('2d');
        new Chart(ctx, {{
            type: 'doughnut',
            data: {{
                labels: ['Critical', 'High', 'Medium', 'Low', 'Informational'],
                datasets: [{{
                    data: [
                        {summary['findings_by_severity']['critical']},
                        {summary['findings_by_severity']['high']},
                        {summary['findings_by_severity']['medium']},
                        {summary['findings_by_severity']['low']},
                        {summary['findings_by_severity']['informational']}
                    ],
                    backgroundColor: ['#d32f2f', '#f57c00', '#fbc02d', '#388e3c', '#0288d1']
                }}]
            }},
            options: {{
                responsive: true,
                maintainAspectRatio: false,
                plugins: {{
                    legend: {{ position: 'bottom' }}
                }}
            }}
        }});
    </script>
</body>
</html>
"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def _generate_alert_html(self, summary: Dict[str, Any]) -> str:
        """Generate alert banner for critical findings."""
        if summary['open_critical_high'] > 0:
            return f'''
        <div class="alert critical">
            <strong>⚠️ ATTENTION REQUIRED:</strong> {summary['open_critical_high']} open critical/high severity findings require immediate action.
        </div>
'''
        return ''
    
    def _get_trend_emoji(self, direction: str) -> str:
        """Get emoji for trend direction."""
        if direction == "improving":
            return "📈✅"
        elif direction == "worsening":
            return "📉⚠️"
        else:
            return "➡️"
    
    def _generate_findings_chart_html(self, summary: Dict[str, Any]) -> str:
        """Generate findings distribution chart section."""
        return f'''
        <div class="section">
            <h2 class="section-title">🔍 Findings Distribution</h2>
            <div class="chart-container">
                <canvas id="findingsChart"></canvas>
            </div>
            <div style="display: grid; grid-template-columns: repeat(5, 1fr); gap: 10px; text-align: center; margin-top: 20px;">
                <div>
                    <div class="badge critical">{summary['findings_by_severity']['critical']}</div>
                    <div style="font-size: 0.85em; margin-top: 5px;">Critical</div>
                </div>
                <div>
                    <div class="badge high">{summary['findings_by_severity']['high']}</div>
                    <div style="font-size: 0.85em; margin-top: 5px;">High</div>
                </div>
                <div>
                    <div class="badge medium">{summary['findings_by_severity']['medium']}</div>
                    <div style="font-size: 0.85em; margin-top: 5px;">Medium</div>
                </div>
                <div>
                    <div class="badge low">{summary['findings_by_severity']['low']}</div>
                    <div style="font-size: 0.85em; margin-top: 5px;">Low</div>
                </div>
                <div>
                    <div style="background: #0288d1; color: white; padding: 4px 12px; border-radius: 12px; font-size: 0.85em; font-weight: 600; display: inline-block;">{summary['findings_by_severity']['informational']}</div>
                    <div style="font-size: 0.85em; margin-top: 5px;">Info</div>
                </div>
            </div>
        </div>
'''
    
    def _generate_compliance_html(self, summary: Dict[str, Any]) -> str:
        """Generate compliance coverage section."""
        return f'''
        <div class="section">
            <h2 class="section-title">✅ Compliance Coverage</h2>
            <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px;">
                {self._compliance_card_html("OWASP Top 10", summary['compliance_coverage']['owasp_coverage'])}
                {self._compliance_card_html("CWE Top 25", summary['compliance_coverage']['cwe_coverage'])}
                {self._compliance_card_html("MITRE ATT&CK", summary['compliance_coverage']['attack_coverage'])}
            </div>
        </div>
'''
    
    def _compliance_card_html(self, name: str, percentage: float) -> str:
        """Generate compliance card HTML."""
        color = "#4caf50" if percentage >= 70 else "#ffc107" if percentage >= 40 else "#f44336"
        return f'''
                <div style="background: white; border: 2px solid #e0e0e0; border-radius: 8px; padding: 20px;">
                    <div style="font-weight: 600; margin-bottom: 10px;">{name}</div>
                    <div style="background: #e0e0e0; height: 20px; border-radius: 10px; overflow: hidden;">
                        <div style="background: {color}; height: 100%; width: {percentage}%; transition: width 0.3s;"></div>
                    </div>
                    <div style="text-align: right; margin-top: 5px; font-weight: bold; color: {color};">{percentage:.1f}%</div>
                </div>
'''
    
    def _generate_top_risks_table_html(self, summary: Dict[str, Any]) -> str:
        """Generate top risky projects table."""
        rows = ""
        for project in summary['top_risky_projects']:
            rows += f'''
                <tr>
                    <td>{project['project_name']}</td>
                    <td><span class="badge {project['color']}">{project['risk_score']}</span></td>
                    <td>{project['total_findings']}</td>
                    <td><span class="badge critical">{project['open_critical_high']}</span></td>
                    <td>{project['severity_counts']['critical']}</td>
                    <td>{project['severity_counts']['high']}</td>
                    <td>{project['severity_counts']['medium']}</td>
                </tr>
'''
        
        return f'''
        <div class="section">
            <h2 class="section-title">🎯 Top Risky Projects</h2>
            <table>
                <thead>
                    <tr>
                        <th>Project</th>
                        <th>Risk Score</th>
                        <th>Total Findings</th>
                        <th>Open Crit/High</th>
                        <th>Critical</th>
                        <th>High</th>
                        <th>Medium</th>
                    </tr>
                </thead>
                <tbody>
                    {rows}
                </tbody>
            </table>
        </div>
'''
    
    def _generate_executive_summary_docx(
        self, 
        summary: Dict[str, Any], 
        projects: List[Project],
        file_path: str
    ):
        """Generate DOCX Executive Summary."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Executive Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Report Generated: {summary['generated_at']}")
        doc.add_paragraph(f"Company: {self.branding.company_name or 'VulnManager'}")
        doc.add_paragraph(f"Coverage: {summary['total_projects']} Active Projects")
        doc.add_paragraph()
        
        # Alert
        if summary['open_critical_high'] > 0:
            p = doc.add_paragraph()
            p.add_run(f"⚠️ ATTENTION: {summary['open_critical_high']} open critical/high findings").bold = True
            doc.add_paragraph()
        
        # KPIs
        doc.add_heading('Key Performance Indicators', 1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        kpis = [
            ('Active Projects', str(summary['total_projects'])),
            ('Total Findings', str(summary['total_findings'])),
            ('MTTR (Days)', str(summary['mttr_days'])),
            ('Trend Direction', summary['trend']['trend_direction'].title()),
            ('Open Critical/High', str(summary['open_critical_high']))
        ]
        
        for i, (label, value) in enumerate(kpis):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Findings by Severity
        doc.add_heading('Findings by Severity', 1)
        severity_table = doc.add_table(rows=6, cols=2)
        severity_table.style = 'Light Grid Accent 1'
        severity_table.rows[0].cells[0].text = 'Severity'
        severity_table.rows[0].cells[1].text = 'Count'
        
        severities = [
            ('Critical', summary['findings_by_severity']['critical']),
            ('High', summary['findings_by_severity']['high']),
            ('Medium', summary['findings_by_severity']['medium']),
            ('Low', summary['findings_by_severity']['low']),
            ('Informational', summary['findings_by_severity']['informational'])
        ]
        
        for i, (sev, count) in enumerate(severities, 1):
            severity_table.rows[i].cells[0].text = sev
            severity_table.rows[i].cells[1].text = str(count)
        
        doc.add_paragraph()
        
        # Compliance Coverage
        doc.add_heading('Compliance Coverage', 1)
        comp_table = doc.add_table(rows=4, cols=2)
        comp_table.style = 'Light Grid Accent 1'
        comp_table.rows[0].cells[0].text = 'Framework'
        comp_table.rows[0].cells[1].text = 'Coverage %'
        
        comp_table.rows[1].cells[0].text = 'OWASP Top 10'
        comp_table.rows[1].cells[1].text = f"{summary['compliance_coverage']['owasp_coverage']:.1f}%"
        comp_table.rows[2].cells[0].text = 'CWE Top 25'
        comp_table.rows[2].cells[1].text = f"{summary['compliance_coverage']['cwe_coverage']:.1f}%"
        comp_table.rows[3].cells[0].text = 'MITRE ATT&CK'
        comp_table.rows[3].cells[1].text = f"{summary['compliance_coverage']['attack_coverage']:.1f}%"
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(self.branding.footer_text)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(file_path)
    
    def _generate_executive_summary_pdf(
        self, 
        summary: Dict[str, Any], 
        projects: List[Project],
        file_path: str
    ):
        """Generate PDF Executive Summary."""
        pdf = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor(self.branding.primary_color),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph("Executive Summary", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        story.append(Paragraph(f"<b>Report Generated:</b> {summary['generated_at']}", styles['Normal']))
        story.append(Paragraph(f"<b>Company:</b> {self.branding.company_name or 'VulnManager'}", styles['Normal']))
        story.append(Paragraph(f"<b>Coverage:</b> {summary['total_projects']} Active Projects", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Alert
        if summary['open_critical_high'] > 0:
            alert_style = ParagraphStyle(
                'Alert',
                parent=styles['Normal'],
                textColor=colors.red,
                fontName='Helvetica-Bold'
            )
            story.append(Paragraph(f"⚠ ATTENTION: {summary['open_critical_high']} open critical/high findings", alert_style))
            story.append(Spacer(1, 0.2*inch))
        
        # KPIs Table
        story.append(Paragraph("Key Performance Indicators", styles['Heading2']))
        kpi_data = [
            ['Metric', 'Value'],
            ['Active Projects', str(summary['total_projects'])],
            ['Total Findings', str(summary['total_findings'])],
            ['MTTR (Days)', str(summary['mttr_days'])],
            ['Trend Direction', summary['trend']['trend_direction'].title()],
            ['Open Critical/High', str(summary['open_critical_high'])]
        ]
        
        kpi_table = Table(kpi_data, colWidths=[3*inch, 2*inch])
        kpi_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.branding.primary_color)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(kpi_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Findings by Severity
        story.append(Paragraph("Findings by Severity", styles['Heading2']))
        severity_data = [
            ['Severity', 'Count'],
            ['Critical', str(summary['findings_by_severity']['critical'])],
            ['High', str(summary['findings_by_severity']['high'])],
            ['Medium', str(summary['findings_by_severity']['medium'])],
            ['Low', str(summary['findings_by_severity']['low'])],
            ['Informational', str(summary['findings_by_severity']['informational'])]
        ]
        
        severity_table = Table(severity_data, colWidths=[3*inch, 2*inch])
        severity_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.branding.primary_color)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(severity_table)
        
        # Footer
        story.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph(self.branding.footer_text, footer_style))
        
        pdf.build(story)
    
    # Placeholder methods for other template types
    def _generate_technical_findings(self, format: ReportFormat, projects: List[Project]) -> str:
        """Generate Technical Findings report with detailed vulnerability information."""
        # Gather all findings from selected projects
        findings_data = []
        total_instances = 0
        
        for project in projects:
            project_findings = self.session.exec(
                select(Finding).where(Finding.project_id == project.id)
            ).all()
            
            for finding in project_findings:
                instances = self.session.exec(
                    select(Instance).where(Instance.finding_id == finding.id)
                ).all()
                
                findings_data.append({
                    'project_name': project.name,
                    'project_id': project.id,
                    'finding': finding,
                    'instances': instances,
                    'instance_count': len(instances)
                })
                total_instances += len(instances)
        
        # Sort by severity (Critical -> High -> Medium -> Low -> Informational)
        severity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3, 'Informational': 4}
        findings_data.sort(key=lambda x: severity_order.get(x['finding'].risk_rating.value, 5))
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == ReportFormat.HTML:
            file_path = f"/tmp/technical_findings_{timestamp}.html"
            self._generate_technical_findings_html(findings_data, projects, total_instances, file_path)
        elif format == ReportFormat.DOCX:
            file_path = f"/tmp/technical_findings_{timestamp}.docx"
            self._generate_technical_findings_docx(findings_data, projects, total_instances, file_path)
        elif format == ReportFormat.PDF:
            file_path = f"/tmp/technical_findings_{timestamp}.pdf"
            self._generate_technical_findings_pdf(findings_data, projects, total_instances, file_path)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        return file_path
    
    def _generate_technical_findings_html(
        self, 
        findings_data: List[Dict], 
        projects: List[Project],
        total_instances: int,
        file_path: str
    ):
        """Generate interactive HTML Technical Findings report."""
        # Count findings by severity
        severity_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0, 'Informational': 0}
        for item in findings_data:
            severity = item['finding'].risk_rating.value
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        # Build findings HTML
        findings_html = ""
        for idx, item in enumerate(findings_data, 1):
            finding = item['finding']
            instances = item['instances']
            
            severity_class = finding.risk_rating.value.lower()
            severity_badge_color = {
                'critical': '#d32f2f',
                'high': '#f57c00',
                'medium': '#fbc02d',
                'low': '#388e3c',
                'informational': '#0288d1'
            }.get(severity_class, '#666')
            
            # Build instances table
            instances_rows = ""
            for inst_idx, instance in enumerate(instances, 1):
                instances_rows += f"""
                <tr>
                    <td>{inst_idx}</td>
                    <td><code>{html.escape(instance.location)}</code></td>
                    <td>{html.escape(instance.status)}</td>
                    <td style="font-size: 0.85em;">{instance.created_at.strftime('%Y-%m-%d %H:%M') if instance.created_at else 'N/A'}</td>
                </tr>
                """
            
            findings_html += f"""
            <div class="finding-card" data-severity="{severity_class}">
                <div class="finding-header">
                    <div style="flex: 1;">
                        <span class="badge" style="background: {severity_badge_color};">{finding.risk_rating.value}</span>
                        <span class="finding-id">#{idx}</span>
                        <strong>{html.escape(finding.title)}</strong>
                        <div class="meta-info">
                            <span>🏢 {html.escape(item['project_name'])}</span>
                            <span>📍 {len(instances)} instance(s)</span>
                            {f'<span>🔖 {html.escape(finding.owasp_category)}</span>' if finding.owasp_category else ''}
                            {f'<span>⚖️ {html.escape(finding.issue_status.value)}</span>' if finding.issue_status else ''}
                        </div>
                    </div>
                    <button class="toggle-btn" onclick="toggleDetails({idx})">▼</button>
                </div>
                <div class="finding-details" id="details-{idx}" style="display: none;">
                    <div class="section">
                        <h4>📝 Description</h4>
                        <div class="description">{html.escape(finding.description)}</div>
                    </div>
                    <div class="section">
                        <h4>🛠️ Remediation</h4>
                        <div class="remediation">{html.escape(finding.remediation)}</div>
                    </div>
                    {f'''
                    <div class="section">
                        <h4>🔍 Affected Locations ({len(instances)})</h4>
                        <table class="instances-table">
                            <thead>
                                <tr>
                                    <th style="width: 60px;">#</th>
                                    <th>Location</th>
                                    <th style="width: 150px;">Status</th>
                                    <th style="width: 150px;">Discovered</th>
                                </tr>
                            </thead>
                            <tbody>
                                {instances_rows}
                            </tbody>
                        </table>
                    </div>
                    ''' if instances else '<p style="color: #666; font-style: italic;">No instances recorded.</p>'}
                </div>
            </div>
            """
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Technical Findings - Security Report</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f5f5f5;
        }}
        .container {{
            max-width: 1400px;
            margin: 0 auto;
            padding: 20px;
            background: white;
        }}
        header {{
            background: {self.branding.primary_color};
            color: white;
            padding: 30px;
            text-align: center;
            margin-bottom: 30px;
        }}
        h1 {{
            font-size: 2.5em;
            margin-bottom: 10px;
        }}
        .subtitle {{
            font-size: 1.2em;
            opacity: 0.9;
        }}
        .meta {{
            background: #f9f9f9;
            padding: 15px;
            border-left: 4px solid {self.branding.primary_color};
            margin-bottom: 30px;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 15px;
            margin-bottom: 30px;
        }}
        .stat-card {{
            background: white;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            padding: 15px;
            text-align: center;
        }}
        .stat-value {{
            font-size: 2.5em;
            font-weight: bold;
            color: {self.branding.primary_color};
        }}
        .stat-label {{
            color: #666;
            font-size: 0.85em;
            text-transform: uppercase;
            margin-top: 5px;
        }}
        .filters {{
            background: #f9f9f9;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
        }}
        .filter-buttons {{
            display: flex;
            gap: 10px;
            flex-wrap: wrap;
        }}
        .filter-btn {{
            padding: 8px 16px;
            border: 2px solid #ddd;
            background: white;
            border-radius: 20px;
            cursor: pointer;
            font-size: 0.9em;
            transition: all 0.2s;
        }}
        .filter-btn:hover {{
            background: #f0f0f0;
        }}
        .filter-btn.active {{
            background: {self.branding.primary_color};
            color: white;
            border-color: {self.branding.primary_color};
        }}
        .finding-card {{
            background: white;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            margin-bottom: 15px;
            overflow: hidden;
            transition: all 0.2s;
        }}
        .finding-card:hover {{
            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        }}
        .finding-header {{
            padding: 20px;
            background: #fafafa;
            display: flex;
            align-items: center;
            cursor: pointer;
        }}
        .finding-header:hover {{
            background: #f0f0f0;
        }}
        .finding-id {{
            font-family: monospace;
            color: #666;
            margin-right: 10px;
        }}
        .meta-info {{
            display: flex;
            gap: 15px;
            margin-top: 8px;
            font-size: 0.85em;
            color: #666;
        }}
        .badge {{
            display: inline-block;
            padding: 4px 12px;
            border-radius: 12px;
            font-size: 0.85em;
            font-weight: 600;
            color: white;
            margin-right: 10px;
        }}
        .toggle-btn {{
            background: none;
            border: none;
            font-size: 1.2em;
            cursor: pointer;
            padding: 5px 10px;
            transition: transform 0.2s;
        }}
        .toggle-btn.open {{
            transform: rotate(180deg);
        }}
        .finding-details {{
            padding: 20px;
            border-top: 1px solid #e0e0e0;
        }}
        .section {{
            margin-bottom: 20px;
        }}
        .section h4 {{
            color: {self.branding.primary_color};
            margin-bottom: 10px;
            font-size: 1.1em;
        }}
        .description, .remediation {{
            background: #f9f9f9;
            padding: 15px;
            border-radius: 6px;
            border-left: 4px solid {self.branding.secondary_color};
            white-space: pre-wrap;
        }}
        .instances-table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 10px;
        }}
        .instances-table th, .instances-table td {{
            padding: 10px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }}
        .instances-table th {{
            background: {self.branding.primary_color};
            color: white;
            font-weight: 600;
        }}
        .instances-table tr:hover {{
            background: #f9f9f9;
        }}
        .instances-table code {{
            background: #f0f0f0;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 0.9em;
        }}
        footer {{
            text-align: center;
            padding: 20px;
            color: #666;
            border-top: 1px solid #ddd;
            margin-top: 40px;
        }}
        @media print {{
            body {{ background: white; }}
            .container {{ box-shadow: none; }}
            .filters {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>🔍 Technical Findings</h1>
            <div class="subtitle">Detailed Vulnerability Analysis</div>
        </header>
        
        <div class="meta">
            <strong>Report Generated:</strong> {get_utc_now().isoformat()}<br>
            <strong>Company:</strong> {html.escape(self.branding.company_name or 'VulnManager')}<br>
            <strong>Coverage:</strong> {len(projects)} Project(s) | {len(findings_data)} Finding(s) | {total_instances} Instance(s)
        </div>
        
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-value">{severity_counts.get('Critical', 0)}</div>
                <div class="stat-label">Critical</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{severity_counts.get('High', 0)}</div>
                <div class="stat-label">High</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{severity_counts.get('Medium', 0)}</div>
                <div class="stat-label">Medium</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{severity_counts.get('Low', 0)}</div>
                <div class="stat-label">Low</div>
            </div>
            <div class="stat-card">
                <div class="stat-value">{severity_counts.get('Informational', 0)}</div>
                <div class="stat-label">Informational</div>
            </div>
        </div>
        
        <div class="filters">
            <strong>Filter by Severity:</strong>
            <div class="filter-buttons">
                <button class="filter-btn active" onclick="filterBySeverity('all')">All</button>
                <button class="filter-btn" onclick="filterBySeverity('critical')">Critical</button>
                <button class="filter-btn" onclick="filterBySeverity('high')">High</button>
                <button class="filter-btn" onclick="filterBySeverity('medium')">Medium</button>
                <button class="filter-btn" onclick="filterBySeverity('low')">Low</button>
                <button class="filter-btn" onclick="filterBySeverity('informational')">Informational</button>
            </div>
        </div>
        
        <div id="findings-container">
            {findings_html if findings_html else '<p style="text-align: center; color: #666; padding: 40px;">No findings to display.</p>'}
        </div>
        
        <footer>
            <p>{html.escape(self.branding.footer_text or 'Confidential - Security Assessment Report')}</p>
            <p style="font-size: 0.85em; margin-top: 10px;">Generated by {html.escape(self.branding.company_name or 'VulnManager')} - Advanced Security Assessment Platform</p>
        </footer>
    </div>
    
    <script>
        function toggleDetails(id) {{
            const details = document.getElementById('details-' + id);
            const btn = event.target;
            if (details.style.display === 'none') {{
                details.style.display = 'block';
                btn.classList.add('open');
            }} else {{
                details.style.display = 'none';
                btn.classList.remove('open');
            }}
        }}
        
        function filterBySeverity(severity) {{
            const cards = document.querySelectorAll('.finding-card');
            const buttons = document.querySelectorAll('.filter-btn');
            
            // Update button states
            buttons.forEach(btn => btn.classList.remove('active'));
            event.target.classList.add('active');
            
            // Filter cards
            cards.forEach(card => {{
                if (severity === 'all' || card.dataset.severity === severity) {{
                    card.style.display = 'block';
                }} else {{
                    card.style.display = 'none';
                }}
            }});
        }}
    </script>
</body>
</html>
"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def _generate_technical_findings_docx(
        self,
        findings_data: List[Dict],
        projects: List[Project],
        total_instances: int,
        file_path: str
    ):
        """Generate DOCX Technical Findings report."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Technical Findings Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Report Generated: {get_utc_now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
        doc.add_paragraph(f"Company: {self.branding.company_name or 'VulnManager'}")
        doc.add_paragraph(f"Coverage: {len(projects)} Project(s) | {len(findings_data)} Finding(s) | {total_instances} Instance(s)")
        doc.add_paragraph("")
        
        # Summary section
        doc.add_heading('Summary', 1)
        severity_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0, 'Informational': 0}
        for item in findings_data:
            severity = item['finding'].risk_rating.value
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        # Summary table
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'Count'
        
        for severity, count in severity_counts.items():
            row_cells = summary_table.add_row().cells
            row_cells[0].text = severity
            row_cells[1].text = str(count)
        
        doc.add_paragraph("")
        
        # Findings section
        doc.add_heading('Findings Details', 1)
        
        for idx, item in enumerate(findings_data, 1):
            finding = item['finding']
            instances = item['instances']
            
            # Finding heading with severity badge
            finding_heading = doc.add_heading(f"Finding #{idx}: {finding.title}", 2)
            
            # Metadata paragraph
            meta_items = [
                f"Severity: {finding.risk_rating.value}",
                f"Project: {item['project_name']}",
                f"Instances: {len(instances)}"
            ]
            if finding.owasp_category:
                meta_items.append(f"OWASP: {finding.owasp_category}")
            if finding.issue_status:
                meta_items.append(f"Status: {finding.issue_status.value}")
            
            doc.add_paragraph(" | ".join(meta_items))
            doc.add_paragraph("")
            
            # Description
            doc.add_heading('Description', 3)
            doc.add_paragraph(finding.description)
            doc.add_paragraph("")
            
            # Remediation
            doc.add_heading('Remediation', 3)
            doc.add_paragraph(finding.remediation)
            doc.add_paragraph("")
            
            # Instances
            if instances:
                doc.add_heading(f'Affected Locations ({len(instances)})', 3)
                instances_table = doc.add_table(rows=1, cols=4)
                instances_table.style = 'Light Grid Accent 1'
                hdr_cells = instances_table.rows[0].cells
                hdr_cells[0].text = '#'
                hdr_cells[1].text = 'Location'
                hdr_cells[2].text = 'Status'
                hdr_cells[3].text = 'Discovered'
                
                for inst_idx, instance in enumerate(instances, 1):
                    row_cells = instances_table.add_row().cells
                    row_cells[0].text = str(inst_idx)
                    row_cells[1].text = instance.location
                    row_cells[2].text = instance.status
                    row_cells[3].text = instance.created_at.strftime('%Y-%m-%d %H:%M') if instance.created_at else 'N/A'
            
            # Page break after each finding (except last)
            if idx < len(findings_data):
                doc.add_page_break()
        
        # Footer
        footer_section = doc.sections[0]
        footer = footer_section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = self.branding.footer_text or 'Confidential - Security Assessment Report'
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(file_path)
    
    def _generate_technical_findings_pdf(
        self,
        findings_data: List[Dict],
        projects: List[Project],
        total_instances: int,
        file_path: str
    ):
        """Generate PDF Technical Findings report."""
        pdf_doc = SimpleDocTemplate(file_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor(self.branding.primary_color),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor(self.branding.primary_color),
            spaceAfter=12
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor(self.branding.secondary_color),
            spaceAfter=6
        )
        
        # Title
        story.append(Paragraph("Technical Findings Report", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Metadata
        story.append(Paragraph(f"<b>Report Generated:</b> {get_utc_now().strftime('%Y-%m-%d %H:%M:%S UTC')}", styles['Normal']))
        story.append(Paragraph(f"<b>Company:</b> {self.branding.company_name or 'VulnManager'}", styles['Normal']))
        story.append(Paragraph(f"<b>Coverage:</b> {len(projects)} Project(s) | {len(findings_data)} Finding(s) | {total_instances} Instance(s)", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Summary
        story.append(Paragraph("Summary", heading_style))
        severity_counts = {'Critical': 0, 'High': 0, 'Medium': 0, 'Low': 0, 'Informational': 0}
        for item in findings_data:
            severity = item['finding'].risk_rating.value
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        summary_data = [['Severity', 'Count']]
        for severity, count in severity_counts.items():
            summary_data.append([severity, str(count)])
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.branding.primary_color)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Findings
        story.append(Paragraph("Findings Details", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        for idx, item in enumerate(findings_data, 1):
            finding = item['finding']
            instances = item['instances']
            
            # Severity color
            severity_colors = {
                'Critical': '#d32f2f',
                'High': '#f57c00',
                'Medium': '#fbc02d',
                'Low': '#388e3c',
                'Informational': '#0288d1'
            }
            severity_color = severity_colors.get(finding.risk_rating.value, '#666666')
            
            # Finding header with colored bar
            story.append(Paragraph(
                f'<b>Finding #{idx}: {finding.title}</b>',
                subheading_style
            ))
            
            # Metadata
            meta_items = [
                f"<b>Severity:</b> {finding.risk_rating.value}",
                f"<b>Project:</b> {item['project_name']}",
                f"<b>Instances:</b> {len(instances)}"
            ]
            if finding.owasp_category:
                meta_items.append(f"<b>OWASP:</b> {finding.owasp_category}")
            if finding.issue_status:
                meta_items.append(f"<b>Status:</b> {finding.issue_status.value}")
            
            story.append(Paragraph(" | ".join(meta_items), styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            
            # Description
            story.append(Paragraph("<b>Description:</b>", styles['Normal']))
            story.append(Paragraph(finding.description, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            
            # Remediation
            story.append(Paragraph("<b>Remediation:</b>", styles['Normal']))
            story.append(Paragraph(finding.remediation, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            
            # Instances
            if instances:
                story.append(Paragraph(f"<b>Affected Locations ({len(instances)}):</b>", styles['Normal']))
                
                instances_data = [['#', 'Location', 'Status', 'Discovered']]
                for inst_idx, instance in enumerate(instances, 1):
                    instances_data.append([
                        str(inst_idx),
                        instance.location[:50] + '...' if len(instance.location) > 50 else instance.location,
                        instance.status,
                        instance.created_at.strftime('%Y-%m-%d') if instance.created_at else 'N/A'
                    ])
                
                instances_table = Table(instances_data, colWidths=[0.5*inch, 3*inch, 1*inch, 1.5*inch])
                instances_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(severity_color)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                ]))
                story.append(instances_table)
            
            # Page break after each finding
            if idx < len(findings_data):
                story.append(PageBreak())
            else:
                story.append(Spacer(1, 0.3*inch))
        
        # Footer
        footer_text = self.branding.footer_text or 'Confidential - Security Assessment Report'
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph(footer_text, styles['Normal']))
        
        pdf_doc.build(story)
    
    def _generate_risk_assessment(self, format: ReportFormat, projects: List[Project]) -> str:
        """Generate Risk Assessment report (placeholder)."""
        return "/tmp/risk_assessment_placeholder.pdf"
    
    def _generate_remediation_status(self, format: ReportFormat, projects: List[Project]) -> str:
        """Generate Remediation Status report (placeholder)."""
        return "/tmp/remediation_status_placeholder.pdf"
    
    def _generate_portfolio_overview(self, format: ReportFormat, projects: List[Project]) -> str:
        """Generate Portfolio Overview report (placeholder)."""
        return "/tmp/portfolio_overview_placeholder.pdf"
