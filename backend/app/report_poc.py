"""DEPRECATED REPORT GENERATOR (Original PoC)

This module is retained only for reference and will be removed in a future
release. The active implementation lives in `report_poc_simple.py` which:

- Uses simplified table structures (no risky merged layouts)
- Generates opaque JPEG donut charts for Word compatibility
- Supports raw and styled variants (/poc, /poc?apply_style=false, /poc/raw)
- Applies left border styling safely (or can skip it)

Why deprecated:
- Original PNG donuts with tight bounding boxes caused Word open errors
- Merged cell approach occasionally produced invalid table XML on some builds
- New sanitization (HTML stripping + Jinja escaping) implemented elsewhere

Do not import this module for new functionality. Prefer `report_poc_simple`.
"""
from __future__ import annotations

from io import BytesIO
from typing import Dict, List

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Cm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# Local models for typing (avoid heavy import cycle)
try:
    from app.models import ProjectReadWithFindings
except Exception:  # pragma: no cover - typing fallback
    ProjectReadWithFindings = object  # type: ignore

# Risk color palette (hex without # for Word XML fill)
RISK_COLORS: Dict[str, str] = {
    "Critical": "8B0000",  # dark red
    "High": "FF4500",      # orange-red
    "Medium": "FFA500",    # orange
    "Low": "9ACD32",       # yellow-green
    "Informational": "1976D2",  # blue (material primary)
}


def _generate_donut_image(label: str, color_hex: str, size_inches: float = 1.5) -> BytesIO:
    """Deprecated PNG donut (kept for reference). Use JPEG variant in simple renderer."""
    fig, ax = plt.subplots(figsize=(size_inches, size_inches))
    ax.pie([1], colors=[f"#{color_hex}"], wedgeprops=dict(width=0.45))
    ax.text(0, 0, label, ha="center", va="center", fontsize=8)
    ax.axis("equal")
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=160, transparent=True, bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)
    buf.seek(0)
    return buf


def _set_table_left_outer_border(tbl, color_hex: str, weight: int = 18) -> None:
    """Apply only the left outer border on a table with the specified color.

    - weight is in eighths of a point; 18 ~ 2.25pt.
    - other borders are set to nil so they remain invisible.
    """
    tblPr = tbl._element.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    # Clear existing
    for child in list(borders):
        borders.remove(child)

    def _border(tag: str, sz: int, val: str, color: str) -> OxmlElement:
        e = OxmlElement(tag)
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        e.set(qn("w:space"), "0")
        return e

    borders.append(_border("w:left", weight, "single", color_hex))
    # Hide other borders
    for side in ("w:top", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        borders.append(_border(side, 0, "nil", "FFFFFF"))


def render_docx_with_template_and_risk_styling(
    template_bytes: bytes,
    project: object,
) -> bytes:
    """Render a report using a provided .docx template then post-process styling.

    Inputs:
    - template_bytes: raw bytes of a .docx file containing Jinja2 placeholders
    - project: ProjectReadWithFindings model

    Output:
    - bytes of a final .docx with left colored border per finding block and
      a donut image reflecting the risk rating.
    """
    # Load template from memory
    tpl = DocxTemplate(BytesIO(template_bytes))

    # Build context for template
    findings_ctx: List[Dict] = []
    for idx, f in enumerate(project.findings, start=1):
        risk = getattr(f, "risk_rating", "Low")
        color = RISK_COLORS.get(risk, "DDDDDD")
        donut_stream = _generate_donut_image("Low" if risk == "Informational" else risk, color)
        
        # Extract instances for affected resources
        instances = getattr(f, "instances", []) or []
        affected_resources = ", ".join([getattr(inst, "location", "") for inst in instances[:3]]) if instances else "N/A"
        if len(instances) > 3:
            affected_resources += f" (and {len(instances) - 3} more)"
        
        # Extract CVE/CWE from description or attributes
        description_text = getattr(f, "description", "") or ""
        import re
        cwe_match = re.search(r'CWE-\d+', description_text, re.IGNORECASE)
        cve_match = re.search(r'CVE-\d{4}-\d+', description_text, re.IGNORECASE)
        cve_cwe = []
        if cve_match:
            cve_cwe.append(cve_match.group(0))
        if cwe_match:
            cve_cwe.append(cwe_match.group(0))
        cve_cwe_text = " / ".join(cve_cwe) if cve_cwe else "N/A"
        
        # Build OWASP vector (if available)
        owasp_category = getattr(f, "owasp_category", None)
        owasp_vector = f"OWASP {owasp_category}" if owasp_category else "N/A"
        
        # Status from instances (first instance status or default)
        status = getattr(instances[0], "status", "New - Unvalidated") if instances else "New - Unvalidated"
        
        findings_ctx.append({
            "index": idx,
            "section_number": f"1.1.{idx}",  # Matches "1.1.1", "1.1.2", etc.
            "risk_rating": risk,
            "risk_color": color,
            "title": getattr(f, "title", "Untitled"),
            "instances_count": len(instances),
            "affected_resources": affected_resources,
            "status": status,
            "cve_cwe": cve_cwe_text,
            "owasp_vector": owasp_vector,
            "impact": "",  # Can be extracted from description or added as field
            # Use plain-text versions to avoid HTML inside docx
            "description_text": description_text,
            "remediation_text": getattr(f, "remediation", "") or "",
            "poc_content": "",  # Placeholder for POC/screenshots
            "donut_img": InlineImage(tpl, donut_stream, Cm(3.5)),  # Larger for left column
        })

    ctx = {
        "project": {"name": getattr(project, "name", "Project")},
        "findings": findings_ctx,
        "total_findings": len(project.findings or []),
    }

    tpl.render(ctx)
    buf = BytesIO()
    tpl.save(buf)
    buf.seek(0)

    # Post-process: set left border color on each finding table
    doc = Document(buf)

    # Heuristic: assume each finding block is a table and the first cell
    # contains either the risk label or we rely on the iteration order to map
    # to findings_ctx. We'll try to detect the risk by reading cell text; if
    # not found, fallback to the i-th color from context.
    fc_iter = iter(findings_ctx)
    for tbl in doc.tables:
        # Try to read risk from the first row's first two cells
        detected_color = None
        try:
            first_text = (tbl.rows[0].cells[0].text or "").strip()
            first_two = first_text + " " + (tbl.rows[0].cells[1].text or "")
            for risk, hexcolor in RISK_COLORS.items():
                if risk.lower() in first_two.lower():
                    detected_color = hexcolor
                    break
        except Exception:
            detected_color = None

        if not detected_color:
            try:
                detected_color = next(fc_iter).get("risk_color")
            except StopIteration:
                detected_color = None

        if detected_color:
            _set_table_left_outer_border(tbl, detected_color)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def build_sample_template_docx() -> bytes:
    """Create a DOCX template matching the user's report structure.

    Layout matches the provided screenshot:
    - Left column: Donut chart (risk colored, circular)
    - Right column: Finding details with labeled rows
    - Structure: Section heading, AFFECTED RESOURCES, STATUS, CVE/CWE, 
      OWASP RISK VECTOR, IMPACT, DESCRIPTION, POC/SCREENSHOT
    - Left border will be colored by risk (applied via post-processing)
    """
    from docx import Document as _WordDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _WordDocument()
    
    # Add a main heading for the findings section
    heading = doc.add_heading("1.1 Critical / Severe Risk Findings", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add spacing
    doc.add_paragraph()

    # Opening loop tag (docxtpl will recognize it and repeat the table)
    doc.add_paragraph("{% for f in findings %}")

    # Create the main finding table: 2 columns, 10 rows
    # Column 0 = narrow for donut, Column 1 = wide for content
    tbl = doc.add_table(rows=10, cols=2)
    tbl.style = "Table Grid"
    
    # Set column widths: left narrow (2"), right wide (4.5")
    tbl.columns[0].width = Inches(2.0)
    tbl.columns[1].width = Inches(4.5)

    # Merge cells in left column (rows 0-9) for donut placement
    cell_left = tbl.cell(0, 0)
    for row_idx in range(1, 10):
        cell_left.merge(tbl.cell(row_idx, 0))
    
    # Place donut image in merged left cell
    cell_left.text = "{{ f.donut_img }}"
    cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # RIGHT COLUMN CONTENT
    # Row 0: Section number + Finding title
    cell_title = tbl.cell(0, 1)
    p_title = cell_title.paragraphs[0]
    p_title.text = "{{ f.section_number }} {{ f.title }}"
    run = p_title.runs[0] if p_title.runs else p_title.add_run()
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Row 1: AFFECTED RESOURCES label + value
    tbl.cell(1, 1).text = "AFFECTED RESOURCES:"
    p1 = tbl.cell(1, 1).paragraphs[0]
    p1.runs[0].font.bold = True
    p1.add_run("\n{{ f.affected_resources }}")
    
    # Row 2: STATUS label + value
    tbl.cell(2, 1).text = "STATUS:"
    p2 = tbl.cell(2, 1).paragraphs[0]
    p2.runs[0].font.bold = True
    p2.add_run("\n{{ f.status }}")
    
    # Row 3: CVE / CWE label + value
    tbl.cell(3, 1).text = "CVE / CWE:"
    p3 = tbl.cell(3, 1).paragraphs[0]
    p3.runs[0].font.bold = True
    p3.add_run("\n{{ f.cve_cwe }}")
    
    # Row 4: OWASP RISK VECTOR label + value
    tbl.cell(4, 1).text = "OWASP RISK VECTOR:"
    p4 = tbl.cell(4, 1).paragraphs[0]
    p4.runs[0].font.bold = True
    p4.add_run("\n{{ f.owasp_vector }}")
    
    # Row 5: IMPACT label + value
    tbl.cell(5, 1).text = "IMPACT:"
    p5 = tbl.cell(5, 1).paragraphs[0]
    p5.runs[0].font.bold = True
    p5.add_run("\n{{ f.impact }}")
    
    # Row 6: DESCRIPTION label
    tbl.cell(6, 1).text = "DESCRIPTION"
    p6 = tbl.cell(6, 1).paragraphs[0]
    p6.runs[0].font.bold = True
    
    # Row 7: Description content (multi-paragraph support)
    tbl.cell(7, 1).text = "{{ f.description_text }}"
    
    # Row 8: POC / SCREENSHOT label
    tbl.cell(8, 1).text = "POC / SCREENSHOT"
    p8 = tbl.cell(8, 1).paragraphs[0]
    p8.runs[0].font.bold = True
    
    # Row 9: POC content (can include images or text)
    tbl.cell(9, 1).text = "{{ f.poc_content }}"

    # Closing loop tag
    doc.add_paragraph("{% endfor %}")
    
    # Add spacing after findings
    doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
