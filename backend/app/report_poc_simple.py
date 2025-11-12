"""
Simplified DOCX report generator without merged cells to avoid corruption.
Uses a side-by-side layout without complex table merging.
"""
from __future__ import annotations

from io import BytesIO
import html
from typing import Dict, List
import re

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Cm, Pt, RGBColor, Inches
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# Risk color palette
RISK_COLORS: Dict[str, str] = {
    "Critical": "8B0000",
    "High": "FF4500",
    "Medium": "FFA500",
    "Low": "9ACD32",
    "Informational": "1976D2",
}

def _normalize_risk_label(risk_val: object) -> str:
    """Return canonical risk label ('Critical', 'High', 'Medium', 'Low', 'Informational').

    Handles enum objects (with .value), strings like 'RiskRating.High', uppercase
    variants ('MEDIUM'), and already-correct values.
    """
    # Enum instance with .value
    try:
        val = getattr(risk_val, "value")
        if isinstance(val, str) and val:
            return val
    except Exception:
        pass

    s = str(risk_val) if risk_val is not None else "Low"
    # Strip enum prefix if present
    if "." in s:
        s = s.split(".")[-1]
    s = s.strip()
    # Normalize common casings
    up = s.upper()
    mapping = {
        "CRITICAL": "Critical",
        "HIGH": "High",
        "MEDIUM": "Medium",
        "LOW": "Low",
        "INFORMATIONAL": "Informational",
        "INFORMATION": "Informational",
        "INFO": "Informational",
    }
    return mapping.get(up, s[:1].upper() + s[1:].lower() if s else "Low")

def _strip_html(text: str) -> str:
    """Remove HTML tags and unescape entities for safe DOCX insertion.

    Word XML cannot contain raw '<' or '>' inside text nodes; ensure we strip tags
    and unescape entities to plain text.
    """
    if not text:
        return ""
    # Remove tags
    txt = re.sub(r"<[^>]+>", "", text)
    # Unescape HTML entities (&amp; &lt; etc.)
    txt = html.unescape(txt)
    return txt


def _fmt_dt(value) -> str:
    """Format datetime/date-like values to ISO date string for templates.

    Returns "N/A" if value is falsy or not datetime-like.
    """
    try:
        if not value:
            return "N/A"
        if hasattr(value, 'strftime'):
            return value.strftime('%Y-%m-%d')
        return str(value)
    except Exception:
        return "N/A"


def _generate_donut_image(label: str, color_hex: str, size_inches: float = 1.5, dpi: int = 150) -> BytesIO:
    """Create a donut chart image with maximum Word compatibility.

    Notes:
    - Use opaque background (no transparency)
    - Save as JPEG to avoid PNG transparency/gamma issues in some Word builds
    """
    fig, ax = plt.subplots(figsize=(size_inches, size_inches))
    ax.pie([1], colors=[f"#{color_hex}"], wedgeprops=dict(width=0.45))
    ax.text(0, 0, label, ha="center", va="center", fontsize=8, fontweight='bold')
    ax.axis("equal")
    buf = BytesIO()
    # Opaque background, standard bounding box to avoid odd drawing anchors
    plt.savefig(buf, format="jpeg", dpi=dpi, facecolor="white", edgecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _set_table_left_border(tbl, color_hex: str, weight: int = 18) -> None:
    """Apply colored left border on table.

    Compatibility: Some Word versions are picky about "nil" border values.
    We set non-left borders to w:val="none" to avoid potential corruption.
    """
    tblPr = tbl._element.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for child in list(borders):
        borders.remove(child)

    def _border(tag: str, sz: int, val: str, color: str) -> OxmlElement:
        e = OxmlElement(tag)
        e.set(qn("w:val"), val)
        # sz and color are ignored when val is "none" but safe to include
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        e.set(qn("w:space"), "0")
        return e

    borders.append(_border("w:left", weight, "single", color_hex))
    # Use val="none" for other borders for maximum Word compatibility
    for side in ("w:top", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        borders.append(_border(side, 0, "none", "auto"))


def render_docx_simple(
    template_bytes: bytes,
    project: object,
    apply_style: bool = True,
    donut_size_cm: float | None = None,
    donut_dpi: int | None = None,
) -> bytes:
    """Render report without complex cell merging.

    Args:
        template_bytes: The raw .docx template content
        project: Object with .name and .findings
        apply_style: If False, skip post-processing (docxtpl-only output)
    """
    tpl = DocxTemplate(BytesIO(template_bytes))

    findings_ctx: List[Dict] = []
    for idx, f in enumerate(project.findings, start=1):
        risk = _normalize_risk_label(getattr(f, "risk_rating", "Low"))
        color = RISK_COLORS.get(risk, "DDDDDD")
        
        try:
            # Determine image size and dpi
            _size_cm = donut_size_cm if donut_size_cm and donut_size_cm > 0 else 2.8
            _dpi = donut_dpi if donut_dpi and donut_dpi > 0 else 150
            _inches = _size_cm / 2.54
            donut_stream = _generate_donut_image(
                risk,
                color,
                size_inches=_inches,
                dpi=_dpi,
            )
            donut_img = InlineImage(tpl, donut_stream, Cm(_size_cm))
        except Exception:
            donut_img = f"[{risk}]"
        
        instances = getattr(f, "instances", []) or []
        affected_resources = ", ".join([getattr(inst, "location", "")[:50] for inst in instances[:3]]) if instances else "N/A"
        if len(instances) > 3:
            affected_resources += f" (+{len(instances) - 3} more)"

        description_text = getattr(f, "description", "") or ""
        description_text = _strip_html(description_text)
        cwe_match = re.search(r'CWE-\d+', description_text, re.IGNORECASE)
        cve_match = re.search(r'CVE-\d{4}-\d+', description_text, re.IGNORECASE)
        cve_cwe_parts = []
        if cve_match:
            cve_cwe_parts.append(cve_match.group(0))
        if cwe_match:
            cve_cwe_parts.append(cwe_match.group(0))
        cve_cwe_text = " / ".join(cve_cwe_parts) if cve_cwe_parts else "N/A"
        
        owasp_category = getattr(f, "owasp_category", None)
        owasp_vector = f"OWASP {owasp_category}" if owasp_category else "N/A"
        status = getattr(instances[0], "status", "New - Unvalidated") if instances else "New - Unvalidated"
        
        findings_ctx.append({
            "index": idx,
            "section_number": f"1.1.{idx}",
            "risk_rating": risk,
            "risk_color": color,
            "title": getattr(f, "title", "Untitled"),
            "instances_count": len(instances),
            "affected_resources": affected_resources,
            "status": status,
            "cve_cwe": cve_cwe_text,
            "owasp_vector": owasp_vector,
            "description_text": description_text[:500] if len(description_text) > 500 else description_text,
            "remediation_text": _strip_html(getattr(f, "remediation", "") or ""),
            "donut_img": donut_img,
            # --- Extended Finding fields exposed as placeholders ---
            "impact": _strip_html(getattr(f, "impact", "") or ""),
            "references_url": getattr(f, "references_url", None) or "N/A",
            "poc_content": _strip_html(getattr(f, "poc_description", "") or ""),
            "review_status": str(getattr(f, "review_status", "Pending") or "Pending"),
            "reviewer_name": getattr(f, "reviewer_name", None) or "N/A",
            "issue_status": str(getattr(f, "issue_status", "Open") or "Open"),
            "issue_status_comment": getattr(f, "issue_status_comment", None) or "",
            "jira_issue_key": getattr(f, "jira_issue_key", None) or "N/A",
            "jira_status": getattr(f, "jira_status", None) or "N/A",
            "remediation_deadline": _fmt_dt(getattr(f, "remediation_deadline", None)),
            "sla_status": str(getattr(f, "sla_status", "") or "N/A"),
            "remediation_owner": getattr(f, "remediation_owner", None) or "N/A",
            "discovered_at": _fmt_dt(getattr(f, "discovered_at", None)),
            "resolved_at": _fmt_dt(getattr(f, "resolved_at", None)),
            "owasp_category": owasp_category or "N/A",
            "cwe_id": getattr(f, "cwe_id", None) or "N/A",
            "cve_id": getattr(f, "cve_id", None) or "N/A",
            "cvss_vector": getattr(f, "cvss_vector", None) or "N/A",
            "cvss_score": getattr(f, "cvss_score", None),
            "owasp_likelihood": getattr(f, "owasp_likelihood", None),
            "owasp_impact": getattr(f, "owasp_impact", None),
            "owasp_risk_rating": getattr(f, "owasp_risk_rating", None) or "N/A",
            "template_id": getattr(f, "template_id", None) or "",
        })

    ctx = {
        "project": {"name": getattr(project, "name", "Project")},
        "findings": findings_ctx,
        "total_findings": len(project.findings or []),
    }

    tpl.render(ctx)
    buf = BytesIO()
    tpl.save(buf)
    buf.seek(0)

    if not apply_style:
        # Return docxtpl-rendered content directly (no post-processing)
        return buf.read()

    # Post-process borders to apply left color per finding table
    doc = Document(buf)
    fc_iter = iter(findings_ctx)

    for tbl in doc.tables:
        try:
            finding_ctx = next(fc_iter)
            color = finding_ctx.get("risk_color")
            if color:
                _set_table_left_border(tbl, color)
        except StopIteration:
            break
        except Exception:
            continue

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def render_docx_raw(template_bytes: bytes, project: object) -> bytes:
    """Render report but replace donut images with plain text and skip all styling.

    Useful for debugging Word open issues by removing drawing relationships and
    any post-processing that edits table XML.
    """
    tpl = DocxTemplate(BytesIO(template_bytes))

    findings_ctx: List[Dict] = []
    for idx, f in enumerate(project.findings, start=1):
        risk = _normalize_risk_label(getattr(f, "risk_rating", "Low"))
        color = RISK_COLORS.get(risk, "DDDDDD")

        # No images: use a short text token instead of InlineImage
        donut_img = f"[{risk}]"

        instances = getattr(f, "instances", []) or []
        affected_resources = ", ".join([getattr(inst, "location", "")[:50] for inst in instances[:3]]) if instances else "N/A"
        if len(instances) > 3:
            affected_resources += f" (+{len(instances) - 3} more)"

        description_text = getattr(f, "description", "") or ""
        description_text = _strip_html(description_text)
        cwe_match = re.search(r'CWE-\d+', description_text, re.IGNORECASE)
        cve_match = re.search(r'CVE-\d{4}-\d+', description_text, re.IGNORECASE)
        cve_cwe_parts = []
        if cve_match:
            cve_cwe_parts.append(cve_match.group(0))
        if cwe_match:
            cve_cwe_parts.append(cwe_match.group(0))
        cve_cwe_text = " / ".join(cve_cwe_parts) if cve_cwe_parts else "N/A"

        owasp_category = getattr(f, "owasp_category", None)
        owasp_vector = f"OWASP {owasp_category}" if owasp_category else "N/A"
        status = getattr(instances[0], "status", "New - Unvalidated") if instances else "New - Unvalidated"

        findings_ctx.append({
            "index": idx,
            "section_number": f"1.1.{idx}",
            "risk_rating": risk,
            "risk_color": color,
            "title": getattr(f, "title", "Untitled"),
            "instances_count": len(instances),
            "affected_resources": affected_resources,
            "status": status,
            "cve_cwe": cve_cwe_text,
            "owasp_vector": owasp_vector,
            "description_text": description_text[:500] if len(description_text) > 500 else description_text,
            "remediation_text": _strip_html(getattr(f, "remediation", "") or ""),
            "donut_img": donut_img,
            # --- Extended Finding fields exposed as placeholders ---
            "impact": _strip_html(getattr(f, "impact", "") or ""),
            "references_url": getattr(f, "references_url", None) or "N/A",
            "poc_content": _strip_html(getattr(f, "poc_description", "") or ""),
            "review_status": str(getattr(f, "review_status", "Pending") or "Pending"),
            "reviewer_name": getattr(f, "reviewer_name", None) or "N/A",
            "issue_status": str(getattr(f, "issue_status", "Open") or "Open"),
            "issue_status_comment": getattr(f, "issue_status_comment", None) or "",
            "jira_issue_key": getattr(f, "jira_issue_key", None) or "N/A",
            "jira_status": getattr(f, "jira_status", None) or "N/A",
            "remediation_deadline": _fmt_dt(getattr(f, "remediation_deadline", None)),
            "sla_status": str(getattr(f, "sla_status", "") or "N/A"),
            "remediation_owner": getattr(f, "remediation_owner", None) or "N/A",
            "discovered_at": _fmt_dt(getattr(f, "discovered_at", None)),
            "resolved_at": _fmt_dt(getattr(f, "resolved_at", None)),
            "owasp_category": owasp_category or "N/A",
            "cwe_id": getattr(f, "cwe_id", None) or "N/A",
            "cve_id": getattr(f, "cve_id", None) or "N/A",
            "cvss_vector": getattr(f, "cvss_vector", None) or "N/A",
            "cvss_score": getattr(f, "cvss_score", None),
            "owasp_likelihood": getattr(f, "owasp_likelihood", None),
            "owasp_impact": getattr(f, "owasp_impact", None),
            "owasp_risk_rating": getattr(f, "owasp_risk_rating", None) or "N/A",
            "template_id": getattr(f, "template_id", None) or "",
        })

    ctx = {
        "project": {"name": getattr(project, "name", "Project")},
        "findings": findings_ctx,
        "total_findings": len(project.findings or []),
    }

    tpl.render(ctx)
    buf = BytesIO()
    tpl.save(buf)
    buf.seek(0)
    return buf.read()


def build_simple_template_docx() -> bytes:
    """Create a simple non-merged template that's less prone to corruption."""
    from docx import Document as _WordDocument
    
    doc = _WordDocument()
    doc.add_heading("1.1 Critical / Severe Risk Findings", level=1)
    doc.add_paragraph()
    doc.add_paragraph("{% for f in findings %}")

    # Simple table: no merged cells
    tbl = doc.add_table(rows=9, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(5.0)

    # Row 0: Donut + Title
    tbl.cell(0, 0).text = "{{ f.donut_img }}"
    tbl.cell(0, 1).text = "{{ f.section_number|e }} {{ f.title|e }}"
    tbl.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(14)
    tbl.cell(0, 1).paragraphs[0].runs[0].font.bold = True

    # Row 1: Affected Resources
    tbl.cell(1, 0).text = "AFFECTED:"
    tbl.cell(1, 1).text = "{{ f.affected_resources|e }}"

    # Row 2: Status
    tbl.cell(2, 0).text = "STATUS:"
    tbl.cell(2, 1).text = "{{ f.status|e }}"

    # Row 3: CVE/CWE
    tbl.cell(3, 0).text = "CVE / CWE:"
    tbl.cell(3, 1).text = "{{ f.cve_cwe|e }}"

    # Row 4: OWASP
    tbl.cell(4, 0).text = "OWASP:"
    tbl.cell(4, 1).text = "{{ f.owasp_vector|e }}"

    # Row 5-6: Description
    tbl.cell(5, 0).text = "DESCRIPTION:"
    tbl.cell(5, 0).paragraphs[0].runs[0].font.bold = True
    tbl.cell(5, 1).text = ""
    tbl.cell(6, 0).text = ""
    tbl.cell(6, 1).text = "{{ f.description_text|e }}"

    # Row 7-8: POC
    tbl.cell(7, 0).text = "POC:"
    tbl.cell(7, 0).paragraphs[0].runs[0].font.bold = True
    tbl.cell(7, 1).text = ""
    tbl.cell(8, 0).text = ""
    tbl.cell(8, 1).text = "(Evidence placeholder)"

    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
