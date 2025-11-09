# backend/app/reports.py

from docx import Document
from docx.shared import Inches
from typing import TYPE_CHECKING, Optional
import os
from datetime import datetime, date
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.units import inch
import html
import re
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend for server
import matplotlib.pyplot as plt
from io import BytesIO

# Type checking import to avoid circular dependencies in SQLModel
if TYPE_CHECKING:
    from app.models import Project # We need the Project model structure for typing

# --- Utility Functions ---

def strip_html_tags(text: str) -> str:
    """
    Remove HTML tags and decode HTML entities from text.
    
    Args:
        text: Text potentially containing HTML markup
        
    Returns:
        Plain text with HTML tags removed and entities decoded
    """
    if not text:
        return ''
    
    # Decode HTML entities first (&nbsp; -> space, &lt; -> <, etc.)
    text = html.unescape(text)
    
    # Remove HTML tags using regex
    # Matches: <tag>, </tag>, <tag attr="value">, etc.
    text = re.sub(r'<[^>]+>', '', text)
    
    # Clean up excess whitespace while preserving paragraph breaks
    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)
    
    # Replace multiple newlines with double newline (paragraph break)
    text = re.sub(r'\n\n+', '\n\n', text)
    
    # Strip leading and trailing whitespace
    text = text.strip()
    
    return text

def generate_report_docx(project: 'Project', file_path: str):
    """
    Generates a security assessment report in DOCX format.
    
    Args:
        project: The Project object including findings and instances.
        file_path: The path where the DOCX file should be saved.
    """
    print(f"Generating DOCX report for project: {project.name}")
    
    # Initialize a new Word document
    document = Document()
    
    # 1. Title Page
    document.add_heading(f"Security Assessment Report: {project.name}", 0)
    document.add_paragraph(f"Consultant: {project.consultant_name or 'N/A'}")
    document.add_paragraph(f"Report Date: {os.path.basename(file_path).split('_')[1].split('.')[0]}")
    document.add_page_break()
    
    # 2. Executive Summary (Placeholder)
    document.add_heading("Executive Summary", 1)
    document.add_paragraph(
        "This is a placeholder for the Executive Summary. The full report details "
        "the findings discovered during the security assessment period."
    )
    
    # 3. Findings Section
    document.add_heading("Detailed Findings", 1)
    
    if not project.findings:
        document.add_paragraph("No findings were recorded for this project.")
    else:
        # Group and sort findings by risk rating (Critical -> High -> Medium, etc.)
        RISK_ORDER = ['Critical', 'High', 'Medium', 'Low', 'Informational']
        sorted_findings = sorted(
            project.findings, 
            key=lambda f: RISK_ORDER.index(f.risk_rating) if f.risk_rating in RISK_ORDER else len(RISK_ORDER)
        )

        for i, finding in enumerate(sorted_findings, 1):
            document.add_heading(f"Finding {i}: {finding.title} ({finding.risk_rating})", 2)
            
            # Risk/Description
            document.add_paragraph(f"Risk Rating: {finding.risk_rating}")
            document.add_heading("Description", 3)
            document.add_paragraph(strip_html_tags(finding.description))
            
            # Remediation
            document.add_heading("Remediation / Solution", 3)
            document.add_paragraph(strip_html_tags(finding.remediation))
            
            # Instances/Locations
            document.add_heading(f"Vulnerable Instances ({len(finding.instances)})", 3)
            
            for instance in finding.instances:
                document.add_paragraph(f"Location: {instance.location}", style='List Bullet')
                document.add_paragraph(f"Details: {strip_html_tags(instance.details)}")
                document.add_paragraph(f"Status: {instance.status}")
                
            document.add_page_break()

    # Save the document
    try:
        document.save(file_path)
        print(f"Report saved to {file_path}")
    except Exception as e:
        print(f"Error saving DOCX file: {e}")
        # Re-raise or handle as needed

def generate_report_pdf(project: 'Project', file_path: str):
    """
    Generates a security assessment report in PDF format using ReportLab.
    
    Args:
        project: The Project object including findings and instances
        file_path: The path where the PDF file should be saved
    """
    print(f"Generating PDF report for project: {project.name}")
    
    # Initialize the PDF document
    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=72,  # 1 inch margins
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30
    )
    heading1_style = styles['Heading1']
    heading2_style = styles['Heading2']
    normal_style = styles['Normal']
    
    # Build the document content
    story = []
    
    # Title Page
    story.append(Paragraph(f"Security Assessment Report", title_style))
    story.append(Paragraph(f"Project: {project.name}", heading1_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Consultant: {project.consultant_name or 'N/A'}", normal_style))
    story.append(Paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}", normal_style))
    story.append(Spacer(1, 60))
    
    # Executive Summary
    story.append(Paragraph("Executive Summary", heading1_style))
    story.append(Paragraph(
        "This security assessment report details the findings discovered during "
        "the assessment period. Each finding is categorized by risk level and "
        "includes detailed information about the vulnerability, its impact, and "
        "recommended remediation steps.",
        normal_style
    ))
    story.append(Spacer(1, 20))
    
    # Risk Summary Table
    risk_counts = {
        'Critical': 0,
        'High': 0,
        'Medium': 0,
        'Low': 0,
        'Informational': 0
    }
    
    for finding in project.findings:
        risk_counts[finding.risk_rating] += 1
    
    # Create risk summary table
    risk_data = [['Risk Level', 'Count']]
    risk_colors = {
        'Critical': colors.red,
        'High': colors.orange,
        'Medium': colors.yellow,
        'Low': colors.green,
        'Informational': colors.blue
    }
    
    for risk in ['Critical', 'High', 'Medium', 'Low', 'Informational']:
        risk_data.append([risk, str(risk_counts[risk])])
    
    risk_table = Table(risk_data, colWidths=[2*inch, inch])
    risk_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 12),
    ]))
    
    story.append(risk_table)
    story.append(Spacer(1, 30))
    
    # Detailed Findings
    story.append(Paragraph("Detailed Findings", heading1_style))
    story.append(Spacer(1, 12))
    
    if not project.findings:
        story.append(Paragraph("No findings were recorded for this project.", normal_style))
    else:
        # Sort findings by risk rating
        RISK_ORDER = ['Critical', 'High', 'Medium', 'Low', 'Informational']
        sorted_findings = sorted(
            project.findings,
            key=lambda f: RISK_ORDER.index(f.risk_rating) if f.risk_rating in RISK_ORDER else len(RISK_ORDER)
        )
        
        for i, finding in enumerate(sorted_findings, 1):
            # Finding Title
            story.append(Paragraph(
                f"Finding {i}: {finding.title}",
                heading2_style
            ))
            
            # Finding details table
            finding_data = [
                ['Risk Rating', finding.risk_rating],
                ['Description', strip_html_tags(finding.description)],
                ['Remediation', strip_html_tags(finding.remediation)],
                ['Instances', str(len(finding.instances))]
            ]
            
            finding_table = Table(finding_data, colWidths=[1.5*inch, 4*inch])
            finding_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ]))
            
            story.append(finding_table)
            story.append(Spacer(1, 12))
            
            # Instance details
            for j, instance in enumerate(finding.instances, 1):
                instance_data = [
                    [f'Instance {j}'],
                    ['Location', instance.location],
                    ['Details', strip_html_tags(instance.details)],
                    ['Status', instance.status]
                ]
                
                instance_table = Table(instance_data, colWidths=[1.5*inch, 4*inch])
                instance_table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('SPAN', (0, 0), (1, 0)),  # Span the title across both columns
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ]))
                
                story.append(instance_table)
                story.append(Spacer(1, 12))
            
            story.append(Spacer(1, 20))
    
    try:
        # Build the PDF
        doc.build(story)
        print(f"PDF report saved to {file_path}")
    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise


def _generate_severity_pie_chart(risk_counts: dict) -> BytesIO:
    """
    Generate a pie chart showing the distribution of findings by severity.
    
    Args:
        risk_counts: Dictionary mapping risk level to count
        
    Returns:
        BytesIO buffer containing the PNG image
    """
    # Define colors matching VulnManager severity scheme
    severity_colors = {
        'Critical': '#d32f2f',  # red
        'High': '#f57c00',      # orange
        'Medium': '#fbc02d',    # yellow
        'Low': '#388e3c',       # green
        'Informational': '#1976d2'  # blue
    }
    
    # Filter out zero counts and prepare data
    labels = []
    sizes = []
    colors_list = []
    
    for risk in ['Critical', 'High', 'Medium', 'Low', 'Informational']:
        count = risk_counts.get(risk, 0)
        if count > 0:
            labels.append(f"{risk} ({count})")
            sizes.append(count)
            colors_list.append(severity_colors[risk])
    
    # Create figure and axis
    fig, ax = plt.subplots(figsize=(8, 6))
    
    # Create pie chart
    wedges, texts, autotexts = ax.pie(
        sizes,
        labels=labels,
        colors=colors_list,
        autopct='%1.1f%%',
        startangle=90,
        textprops={'fontsize': 12}
    )
    
    # Equal aspect ratio ensures that pie is drawn as a circle
    ax.axis('equal')
    
    # Title
    ax.set_title('Findings by Severity', fontsize=16, fontweight='bold', pad=20)
    
    # Save to BytesIO buffer
    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    
    return buf


def _generate_trend_line_chart(project: 'Project', date_from: Optional[date] = None, date_to: Optional[date] = None) -> Optional[BytesIO]:
    """
    Generate a line chart showing the trend of findings over time.
    
    Args:
        project: The Project object with findings
        date_from: Start date for trend analysis (optional)
        date_to: End date for trend analysis (optional)
        
    Returns:
        BytesIO buffer containing the PNG image, or None if insufficient data
    """
    # Extract finding dates (use created_at or discovered_date if available)
    finding_dates = []
    for finding in project.findings:
        if hasattr(finding, 'created_at') and finding.created_at:
            finding_dates.append(finding.created_at.date())
    
    # Need at least 2 data points for a trend
    if len(finding_dates) < 2:
        return None
    
    # Sort dates
    finding_dates.sort()
    
    # Apply date filters if provided
    if date_from:
        finding_dates = [d for d in finding_dates if d >= date_from]
    if date_to:
        finding_dates = [d for d in finding_dates if d <= date_to]
    
    if len(finding_dates) < 2:
        return None
    
    # Count findings per date
    from collections import Counter
    date_counts = Counter(finding_dates)
    
    # Calculate cumulative counts
    sorted_dates = sorted(date_counts.keys())
    cumulative_counts = []
    total = 0
    
    for d in sorted_dates:
        total += date_counts[d]
        cumulative_counts.append(total)
    
    # Create figure and axis
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Plot cumulative trend
    ax.plot(sorted_dates, cumulative_counts, marker='o', linestyle='-', 
            linewidth=2, color='#1976d2', markersize=6)
    
    # Fill area under the curve
    ax.fill_between(sorted_dates, cumulative_counts, alpha=0.3, color='#1976d2')
    
    # Labels and title
    ax.set_xlabel('Date', fontsize=12, fontweight='bold')
    ax.set_ylabel('Cumulative Findings', fontsize=12, fontweight='bold')
    ax.set_title('Findings Discovery Trend', fontsize=16, fontweight='bold', pad=20)
    
    # Grid for better readability
    ax.grid(True, alpha=0.3, linestyle='--')
    
    # Rotate x-axis labels for better fit
    plt.xticks(rotation=45, ha='right')
    
    # Save to BytesIO buffer
    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    
    return buf


def generate_executive_report_pdf(
    project: 'Project',
    file_path: str,
    date_from: Optional[date] = None,
    date_to: Optional[date] = None,
    include_charts: bool = True,
    logo_url: Optional[str] = None,
    company_name: Optional[str] = None,
    custom_header: Optional[str] = None,
    custom_footer: Optional[str] = None
):
    """
    Generates an executive summary report in PDF format with charts and branding.
    
    This is a more polished, stakeholder-friendly report compared to the detailed
    technical report. It focuses on high-level metrics, visualizations, and
    actionable insights.
    
    Args:
        project: The Project object including findings and instances
        file_path: The path where the PDF file should be saved
        date_from: Optional start date for filtering findings
        date_to: Optional end date for filtering findings
        include_charts: Whether to include charts in the report
        logo_url: Optional URL or path to company logo
        company_name: Optional company name for branding
        custom_header: Optional custom header text
        custom_footer: Optional custom footer text
    """
    print(f"Generating Executive Report PDF for project: {project.name}")
    
    # Initialize the PDF document
    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles for executive report
    title_style = ParagraphStyle(
        'ExecutiveTitle',
        parent=styles['Heading1'],
        fontSize=28,
        spaceAfter=30,
        textColor=colors.HexColor('#1976d2'),
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'ExecutiveSubtitle',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=20,
        textColor=colors.HexColor('#455a64'),
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'ExecutiveHeading',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=12,
        textColor=colors.HexColor('#1976d2'),
        fontName='Helvetica-Bold'
    )
    
    normal_style = styles['Normal']
    
    # Build the document content
    story = []
    
    # === TITLE PAGE ===
    story.append(Spacer(1, 30))
    
    # Company logo (if provided)
    # TODO(future): Load logo from branding settings (requires logo upload feature)
    
    story.append(Paragraph("Executive Security Report", title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(project.name, subtitle_style))
    story.append(Spacer(1, 40))
    
    # Metadata table
    report_date = datetime.now().strftime('%B %d, %Y')
    metadata = [
        ['Report Date:', report_date],
        ['Consultant:', project.consultant_name or 'N/A'],
        ['Project ID:', str(project.id)],
    ]
    
    if company_name:
        metadata.append(['Company:', company_name])
    
    if date_from and date_to:
        metadata.append(['Assessment Period:', f"{date_from} to {date_to}"])
    
    metadata_table = Table(metadata, colWidths=[2*inch, 3.5*inch])
    metadata_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    
    story.append(metadata_table)
    story.append(Spacer(1, 30))
    
    # Custom header text (if provided)
    if custom_header:
        story.append(Paragraph(custom_header, normal_style))
        story.append(Spacer(1, 20))
    
    story.append(PageBreak())
    
    # === EXECUTIVE SUMMARY ===
    story.append(Paragraph("Executive Summary", subtitle_style))
    story.append(Spacer(1, 12))
    
    # Calculate risk metrics
    total_findings = len(project.findings)
    risk_counts = {
        'Critical': 0,
        'High': 0,
        'Medium': 0,
        'Low': 0,
        'Informational': 0
    }
    
    for finding in project.findings:
        risk_counts[finding.risk_rating] = risk_counts.get(finding.risk_rating, 0) + 1
    
    # Executive summary text
    critical_high_count = risk_counts['Critical'] + risk_counts['High']
    
    summary_text = f"""
    This executive report provides a high-level overview of the security assessment 
    conducted for {project.name}. During the assessment period, {total_findings} security 
    findings were identified, of which {critical_high_count} are classified as 
    Critical or High severity and require immediate attention.
    """
    
    story.append(Paragraph(summary_text.strip(), normal_style))
    story.append(Spacer(1, 20))
    
    # === KEY METRICS ===
    story.append(Paragraph("Key Metrics", heading_style))
    story.append(Spacer(1, 12))
    
    # Create metrics summary table
    metrics_data = [
        ['Total Findings', str(total_findings)],
        ['Critical Severity', str(risk_counts['Critical'])],
        ['High Severity', str(risk_counts['High'])],
        ['Medium Severity', str(risk_counts['Medium'])],
        ['Low Severity', str(risk_counts['Low'])],
        ['Informational', str(risk_counts['Informational'])],
    ]
    
    metrics_table = Table(metrics_data, colWidths=[3*inch, 1.5*inch])
    metrics_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e3f2fd')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        # Highlight critical and high counts
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#ffebee')),  # Critical
        ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor('#fff3e0')),  # High
    ]))
    
    story.append(metrics_table)
    story.append(Spacer(1, 30))
    
    # === CHARTS ===
    if include_charts and total_findings > 0:
        # Severity distribution pie chart
        pie_chart_buf = _generate_severity_pie_chart(risk_counts)
        if pie_chart_buf:
            chart_img = Image(pie_chart_buf, width=5*inch, height=3.75*inch)
            story.append(chart_img)
            story.append(Spacer(1, 20))
        
        # Trend line chart (if enough data)
        trend_chart_buf = _generate_trend_line_chart(project, date_from, date_to)
        if trend_chart_buf:
            story.append(PageBreak())
            story.append(Paragraph("Findings Discovery Trend", heading_style))
            story.append(Spacer(1, 12))
            chart_img = Image(trend_chart_buf, width=6*inch, height=3.6*inch)
            story.append(chart_img)
            story.append(Spacer(1, 20))
    
    story.append(PageBreak())
    
    # === TOP 5 CRITICAL FINDINGS ===
    story.append(Paragraph("Top Priority Findings", heading_style))
    story.append(Spacer(1, 12))
    
    # Sort findings by risk (Critical > High > Medium > Low > Informational)
    RISK_ORDER = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    sorted_findings = sorted(
        project.findings,
        key=lambda f: RISK_ORDER.index(f.risk_rating) if f.risk_rating in RISK_ORDER else len(RISK_ORDER)
    )
    
    # Take top 5 findings
    top_findings = sorted_findings[:5]
    
    if not top_findings:
        story.append(Paragraph("No findings were recorded for this project.", normal_style))
    else:
        # Create top findings table
        top_findings_data = [['#', 'Finding', 'Severity', 'Instances']]
        
        for i, finding in enumerate(top_findings, 1):
            top_findings_data.append([
                str(i),
                finding.title[:60] + ('...' if len(finding.title) > 60 else ''),
                finding.risk_rating,
                str(len(finding.instances))
            ])
        
        top_table = Table(top_findings_data, colWidths=[0.4*inch, 3.2*inch, 1.2*inch, 0.8*inch])
        top_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
            ('ALIGN', (3, 0), (3, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f5f5f5')),
        ]))
        
        story.append(top_table)
        story.append(Spacer(1, 20))
    
    # === RECOMMENDATIONS ===
    story.append(PageBreak())
    story.append(Paragraph("Recommendations", heading_style))
    story.append(Spacer(1, 12))
    
    recommendations = []
    
    if risk_counts['Critical'] > 0:
        recommendations.append(
            f"<b>Critical Priority:</b> Address all {risk_counts['Critical']} Critical severity "
            f"findings immediately. These vulnerabilities pose a direct threat to system security "
            f"and should be remediated within 24-48 hours."
        )
    
    if risk_counts['High'] > 0:
        recommendations.append(
            f"<b>High Priority:</b> Remediate {risk_counts['High']} High severity findings within "
            f"1-2 weeks. These issues present significant security risks and should be prioritized "
            f"in the remediation plan."
        )
    
    if critical_high_count == 0 and risk_counts['Medium'] > 0:
        recommendations.append(
            "<b>Good Security Posture:</b> No Critical or High severity findings were identified. "
            "Focus on addressing Medium severity issues to further strengthen security controls."
        )
    
    recommendations.append(
        "<b>Ongoing Assessment:</b> Conduct regular security assessments (quarterly recommended) "
        "to identify new vulnerabilities and validate remediation efforts."
    )
    
    recommendations.append(
        "<b>Security Training:</b> Provide security awareness training to development teams to "
        "prevent common vulnerabilities from being introduced during the development lifecycle."
    )
    
    for rec in recommendations:
        story.append(Paragraph(rec, normal_style))
        story.append(Spacer(1, 12))
    
    # Custom footer (if provided)
    if custom_footer:
        story.append(Spacer(1, 30))
        story.append(Paragraph(custom_footer, normal_style))
    
    # Build the PDF
    try:
        doc.build(story)
        print(f"Executive Report PDF saved to {file_path}")
    except Exception as e:
        print(f"Error generating executive report PDF: {e}")
        raise

# --- Template Rendering Engine (v1.1.0) ---

def render_template(
    template,
    project,
    file_path: str,
    variables: Optional[dict] = None
) -> None:
    """
    Render a report from a template.
    
    This function processes the template's sections configuration and generates
    a PDF report based on the enabled sections and their settings.
    
    Args:
        template: The ReportTemplate object with sections and variables config
        project: The Project object including findings and instances
        file_path: The path where the PDF file should be saved
        variables: Dict of variable values to use (overrides template defaults)
    """
    import json
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    
    print(f"Rendering report from template: {template.name} for project: {project.name}")
    
    try:
        # Parse template configuration
        sections = json.loads(template.sections) if template.sections else []
        template_vars = json.loads(template.variables) if template.variables else []
        
        # Build variable defaults dict
        var_defaults = {var['name']: var.get('default') for var in template_vars}
        
        # Merge with provided variables
        render_vars = {**var_defaults, **(variables or {})}
        
        # Sort sections by order
        sections_sorted = sorted(
            [s for s in sections if s.get('enabled', True)],
            key=lambda x: x.get('order', 999)
        )
        
        # Initialize PDF
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Process each section
        for section in sections_sorted:
            section_id = section.get('id')
            section_settings = section.get('settings', {})
            
            if section_id == 'title':
                # Title page section
                company_name = render_vars.get('company_name', '')
                
                story.append(Spacer(1, 2*inch))
                
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Title'],
                    fontSize=28,
                    textColor=colors.HexColor('#1976d2'),
                    spaceAfter=30,
                    alignment=1  # Center
                )
                story.append(Paragraph(f"Security Assessment Report", title_style))
                story.append(Paragraph(f"<b>{project.name}</b>", styles['Title']))
                
                if company_name:
                    story.append(Spacer(1, 0.5*inch))
                    story.append(Paragraph(f"Prepared for: {company_name}", styles['Normal']))
                
                story.append(Spacer(1, 0.5*inch))
                story.append(Paragraph(
                    f"Report Date: {datetime.now().strftime('%B %d, %Y')}",
                    styles['Normal']
                ))
                
                if project.consultant_name:
                    story.append(Paragraph(f"Consultant: {project.consultant_name}", styles['Normal']))
                
                story.append(PageBreak())
            
            elif section_id == 'summary':
                # Executive summary section
                story.append(Paragraph("Executive Summary", styles['Heading1']))
                story.append(Spacer(1, 0.2*inch))
                
                # Calculate metrics
                total_findings = len(project.findings) if project.findings else 0
                risk_counts = {}
                if project.findings:
                    for finding in project.findings:
                        risk = finding.risk_rating
                        risk_counts[risk] = risk_counts.get(risk, 0) + 1
                
                # Summary text
                summary_text = (
                    f"This security assessment of <b>{project.name}</b> identified "
                    f"<b>{total_findings}</b> finding{'s' if total_findings != 1 else ''} "
                    f"across various risk levels."
                )
                story.append(Paragraph(summary_text, styles['Normal']))
                story.append(Spacer(1, 0.3*inch))
            
            elif section_id == 'charts':
                # Charts section
                include_pie = section_settings.get('include_pie', True) and render_vars.get('include_charts', True)
                include_line = section_settings.get('include_line', True) and render_vars.get('include_charts', True)
                
                if include_pie or include_line:
                    story.append(Paragraph("Risk Analysis", styles['Heading1']))
                    story.append(Spacer(1, 0.2*inch))
                
                if include_pie and project.findings:
                    # Generate pie chart
                    risk_counts = {}
                    for finding in project.findings:
                        risk = finding.risk_rating
                        risk_counts[risk] = risk_counts.get(risk, 0) + 1
                    
                    chart_buffer = _generate_severity_pie_chart(risk_counts)
                    if chart_buffer:
                        img = Image(chart_buffer, width=6*inch, height=4*inch)
                        story.append(img)
                        story.append(Spacer(1, 0.3*inch))
                
                if include_line and project.findings:
                    # Generate trend chart
                    chart_buffer = _generate_trend_line_chart(project, None, None)
                    if chart_buffer:
                        img = Image(chart_buffer, width=6*inch, height=4*inch)
                        story.append(img)
                        story.append(Spacer(1, 0.3*inch))
            
            elif section_id == 'findings':
                # Top findings section
                max_items = section_settings.get('max_items', render_vars.get('max_findings', 10))
                group_by = section_settings.get('group_by', 'risk')
                
                story.append(Paragraph(f"Top {max_items} Findings", styles['Heading1']))
                story.append(Spacer(1, 0.2*inch))
                
                if project.findings:
                    # Sort by risk rating
                    RISK_ORDER = ['Critical', 'High', 'Medium', 'Low', 'Informational']
                    sorted_findings = sorted(
                        project.findings,
                        key=lambda f: RISK_ORDER.index(f.risk_rating) if f.risk_rating in RISK_ORDER else len(RISK_ORDER)
                    )[:max_items]
                    
                    for i, finding in enumerate(sorted_findings, 1):
                        # Finding title
                        finding_title = f"{i}. {finding.title} ({finding.risk_rating})"
                        story.append(Paragraph(finding_title, styles['Heading2']))
                        
                        # Description
                        desc_clean = strip_html_tags(finding.description)
                        if len(desc_clean) > 300:
                            desc_clean = desc_clean[:300] + "..."
                        story.append(Paragraph(desc_clean, styles['Normal']))
                        story.append(Spacer(1, 0.2*inch))
                else:
                    story.append(Paragraph("No findings recorded.", styles['Normal']))
            
            elif section_id == 'recommendations':
                # Recommendations section
                story.append(Paragraph("Recommendations", styles['Heading1']))
                story.append(Spacer(1, 0.2*inch))
                
                # Custom footer if provided
                custom_footer = render_vars.get('custom_footer', '')
                if custom_footer:
                    story.append(Paragraph(custom_footer, styles['Normal']))
                else:
                    # Default recommendations based on risk profile
                    if project.findings:
                        critical_count = sum(1 for f in project.findings if f.risk_rating == 'Critical')
                        high_count = sum(1 for f in project.findings if f.risk_rating == 'High')
                        
                        if critical_count > 0:
                            story.append(Paragraph(
                                f"• Immediate action required: Address {critical_count} critical finding(s)",
                                styles['Normal']
                            ))
                        if high_count > 0:
                            story.append(Paragraph(
                                f"• High priority: Remediate {high_count} high-risk finding(s) within 30 days",
                                styles['Normal']
                            ))
                        story.append(Paragraph(
                            "• Implement regular security assessments",
                            styles['Normal']
                        ))
                    else:
                        story.append(Paragraph("Continue maintaining security best practices.", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        print(f"Template report saved to {file_path}")
        
    except Exception as e:
        print(f"Error rendering template report: {e}")
        raise
