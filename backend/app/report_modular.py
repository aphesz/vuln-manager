"""
Modular DOCX report generation system.

Supports user-uploaded custom DOCX templates with placeholders (like report_poc_simple.py).
Templates are stored in database and loaded from storage/templates/.

v0.12.0: Unified template system - users can upload custom DOCX templates
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import List, Dict, Optional, Any, Union, Set
import json
import re
from datetime import datetime

from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Cm
from docxcompose.composer import Composer
from sqlmodel import Session, select

from app.models import ReportTemplate
from app.report_poc_simple import (
    _strip_html, 
    _fmt_dt, 
    _normalize_risk_label,
    _generate_donut_image,
    _set_table_left_border,
    RISK_COLORS,
)


# Storage root for templates
# Use /code/storage in container, fallback to relative path for local dev
STORAGE_ROOT = Path("/code/storage/templates") if Path("/code/storage").exists() else Path(__file__).parent.parent / "storage" / "templates"

# Legacy module directory (for backward compatibility)
MODULE_DIR = Path(__file__).parent / "report_modules"


def generate_sample_project_data(project_name: str = "Sample Security Assessment") -> Dict[str, Any]:
    """Generate realistic sample data for template preview.
    
    Creates fake findings, risk counts, and project metadata for testing templates
    without requiring real database data.
    
    Args:
        project_name: Name for the sample project
        
    Returns:
        Complete context dict compatible with build_context()
    """
    sample_findings = [
        {
            "title": "Cross-Site Scripting (XSS) in Search Parameter",
            "risk_rating": "Critical",
            "description_text": "The application fails to properly sanitize user input in the search parameter, "
                              "allowing attackers to inject arbitrary JavaScript code. This vulnerability affects "
                              "all authenticated users and can lead to session hijacking, credential theft, and "
                              "unauthorized actions performed on behalf of victims.",
            "impact": "Attackers can execute malicious scripts in victim browsers, steal session tokens, "
                     "redirect users to phishing sites, or perform actions as the authenticated user.",
            "remediation_text": "Implement proper input validation and output encoding. Use a Content Security Policy (CSP) "
                               "to restrict inline script execution. Encode all user-controllable data before rendering in HTML context.",
            "poc_content": "1. Navigate to https://example.com/search?q=<script>alert(document.cookie)</script>\n"
                          "2. Observe the JavaScript execution in the browser\n"
                          "3. Cookie data is displayed in the alert dialog",
            "affected_resources": "https://example.com/search, https://example.com/products/search (+2 more)",
            "instances_count": 4,
            "cve_id": "N/A",
            "cwe_id": "CWE-79",
            "cvss_score": 9.3,
            "cvss_vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:H/I:H/A:N",
            "owasp_category": "A03:2021",
            "owasp_risk_rating": "Critical",
        },
        {
            "title": "SQL Injection in User Profile Endpoint",
            "risk_rating": "High",
            "description_text": "The user profile API endpoint concatenates user input directly into SQL queries "
                              "without proper parameterization. An attacker can manipulate the 'user_id' parameter "
                              "to execute arbitrary SQL commands, potentially accessing or modifying database contents.",
            "impact": "Complete database compromise including unauthorized access to sensitive user data, "
                     "modification of records, and potential server takeover through advanced SQL injection techniques.",
            "remediation_text": "Use prepared statements with parameterized queries for all database interactions. "
                               "Implement least-privilege database access. Deploy web application firewall (WAF) rules.",
            "poc_content": "POST /api/user/profile\nContent-Type: application/json\n\n"
                          "{\"user_id\": \"1' OR '1'='1\"}",
            "affected_resources": "https://api.example.com/user/profile",
            "instances_count": 1,
            "cve_id": "N/A",
            "cwe_id": "CWE-89",
            "cvss_score": 8.6,
            "cvss_vector": "CVSS:3.1/AV:N/AC:L/PR:L/UI:N/S:U/C:H/I:H/A:L",
            "owasp_category": "A03:2021",
            "owasp_risk_rating": "High",
        },
        {
            "title": "Broken Authentication - Weak Password Policy",
            "risk_rating": "Medium",
            "description_text": "The application allows users to create accounts with weak passwords (minimum 4 characters, "
                              "no complexity requirements). This weakness enables brute-force and dictionary attacks.",
            "impact": "User accounts can be compromised through automated password guessing attacks, "
                     "leading to unauthorized access and potential data breaches.",
            "remediation_text": "Enforce strong password policy (minimum 12 characters, complexity requirements). "
                               "Implement account lockout after failed attempts. Use multi-factor authentication (MFA).",
            "poc_content": "Successfully created account with password '1234' during testing.",
            "affected_resources": "https://example.com/register",
            "instances_count": 1,
            "cve_id": "N/A",
            "cwe_id": "CWE-521",
            "cvss_score": 5.3,
            "cvss_vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N",
            "owasp_category": "A07:2021",
            "owasp_risk_rating": "Medium",
        },
    ]
    
    # Add section numbers and other required fields
    for idx, finding in enumerate(sample_findings, start=1):
        finding.update({
            "index": idx,
            "section_number": f"1.1.{idx}",
            "references_url": "https://owasp.org/www-project-top-ten/",
            "issue_status": "Open",
            "review_status": "Pending",
            "reviewer_name": "N/A",
            "sla_status": "At Risk" if idx <= 2 else "On Track",
            "remediation_deadline": "2025-12-31",
            "remediation_owner": "Development Team",
            "jira_issue_key": f"VULN-{100 + idx}",
            "jira_status": "To Do",
            "discovered_at": "2025-11-01",
            "resolved_at": "N/A",
            "template_id": "",
            "status": "New - Unvalidated",
            "owasp_vector": f"OWASP {finding['owasp_category']}",
            "owasp_likelihood": None,
            "owasp_impact": None,
        })
    
    return {
        "project": {
            "name": project_name,
            "consultant_name": "John Security Analyst",
        },
        "findings": sample_findings,
        "total_findings": len(sample_findings),
        "critical_count": 1,
        "high_count": 1,
        "medium_count": 1,
        "low_count": 0,
        "informational_count": 0,
        "overdue_count": 0,
        "at_risk_count": 2,
        "on_track_count": 1,
        "company_name": "ACME Corporation",
        "report_date": datetime.now().strftime("%Y-%m-%d"),
        "report_version": "1.0 - PREVIEW",
        "consultant_email": "consultant@example.com",
        "assessment_period": "Q4 2025",
    }


def extract_jinja2_variables(docx_path: Path) -> List[Dict[str, Any]]:
    """Extract Jinja2 variables from a DOCX template file.
    
    Parses the document XML to find all Jinja2 syntax:
    - {{ variable }} - simple variable interpolation
    - {% for item in items %} - loop variables
    - {% if condition %} - conditional variables
    
    Args:
        docx_path: Path to DOCX template file
        
    Returns:
        List of variable dictionaries with metadata:
        [
            {
                "name": "project_name",
                "type": "string",
                "required": True,
                "context": "simple",
                "sample_value": ""
            },
            {
                "name": "findings",
                "type": "list",
                "required": True,
                "context": "loop",
                "sample_value": []
            }
        ]
    """
    doc = Document(docx_path)
    variables: Dict[str, Dict[str, Any]] = {}
    
    # Regex patterns for Jinja2 syntax
    variable_pattern = re.compile(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_\.]*)\s*(?:\|[^}]*)?\}\}')
    for_pattern = re.compile(r'\{%\s*for\s+([a-zA-Z_][a-zA-Z0-9_]*)\s+in\s+([a-zA-Z_][a-zA-Z0-9_\.]*)\s*%\}')
    if_pattern = re.compile(r'\{%\s*if\s+([a-zA-Z_][a-zA-Z0-9_\.]*)\s*(?:[^%]*)%\}')
    
    # Extract text from all paragraphs and tables
    all_text = []
    for paragraph in doc.paragraphs:
        all_text.append(paragraph.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_text.append(paragraph.text)
    
    # Also check headers/footers
    for section in doc.sections:
        for header in [section.header, section.footer]:
            for paragraph in header.paragraphs:
                all_text.append(paragraph.text)
    
    full_text = '\n'.join(all_text)
    
    # Find simple variables {{ var }}
    for match in variable_pattern.finditer(full_text):
        var_name = match.group(1)
        # Skip nested properties for now, just get root
        root_var = var_name.split('.')[0]
        
        if root_var not in variables:
            # Infer type from name patterns
            inferred_type = "string"
            if root_var.endswith('_count') or root_var.endswith('_score'):
                inferred_type = "number"
            elif root_var.endswith('_date') or root_var.startswith('date_'):
                inferred_type = "date"
            elif root_var.endswith('s') and not root_var.endswith('ss'):  # plural
                inferred_type = "list"
            
            variables[root_var] = {
                "name": root_var,
                "type": inferred_type,
                "required": True,
                "context": "simple",
                "sample_value": _get_sample_value(inferred_type)
            }
    
    # Find loop variables {% for item in items %}
    for match in for_pattern.finditer(full_text):
        loop_var = match.group(1)  # item
        collection_var = match.group(2)  # items
        
        if collection_var not in variables:
            variables[collection_var] = {
                "name": collection_var,
                "type": "list",
                "required": True,
                "context": "loop",
                "sample_value": []
            }
    
    # Find conditional variables {% if var %}
    for match in if_pattern.finditer(full_text):
        var_name = match.group(1)
        root_var = var_name.split('.')[0]
        
        if root_var not in variables:
            variables[root_var] = {
                "name": root_var,
                "type": "boolean",
                "required": False,  # conditionals are optional
                "context": "conditional",
                "sample_value": False
            }
    
    return list(variables.values())


def _get_sample_value(var_type: str) -> Any:
    """Get sample value for a variable type."""
    samples = {
        "string": "",
        "number": 0,
        "boolean": False,
        "date": datetime.now().strftime("%Y-%m-%d"),
        "list": []
    }
    return samples.get(var_type, "")


def get_template_by_id(session: Session, template_id: int) -> Optional[ReportTemplate]:
    """Load template from database by ID.
    
    Args:
        session: Database session
        template_id: Template ID
        
    Returns:
        ReportTemplate object or None
    """
    return session.get(ReportTemplate, template_id)


def get_template_by_name(session: Session, name: str, project_id: Optional[int] = None) -> Optional[ReportTemplate]:
    """Load template from database by name.
    
    Args:
        session: Database session  
        name: Template name
        project_id: Optional project ID for project-specific templates
        
    Returns:
        ReportTemplate object or None
    """
    # Try project-specific first
    if project_id:
        stmt = select(ReportTemplate).where(
            ReportTemplate.name == name,
            ReportTemplate.project_id == project_id
        )
        result = session.exec(stmt).first()
        if result:
            return result
    
    # Fall back to system/public templates
    stmt = select(ReportTemplate).where(
        ReportTemplate.name == name,
        ReportTemplate.is_system_template == True
    )
    return session.exec(stmt).first()


def get_template_path(template: ReportTemplate) -> Path:
    """Get filesystem path for a template.
    
    Args:
        template: ReportTemplate object
        
    Returns:
        Path to DOCX file
        
    Raises:
        FileNotFoundError: If template file doesn't exist
    """
    # DOCX templates stored in storage/templates/
    if template.docx_file_path:
        path = STORAGE_ROOT / template.docx_file_path
        if not path.exists():
            raise FileNotFoundError(f"Template file not found: {path}")
        return path
    
    # Legacy fallback: should not happen with new system
    raise ValueError(f"Template '{template.name}' has no docx_file_path")


def get_module_path(module_name: str) -> Path:
    """DEPRECATED: Get the file path for a module template.
    
    This function is kept for backward compatibility.
    New code should use get_template_by_id() or get_template_by_name().
    
    Args:
        module_name: Module name without .docx extension
        
    Returns:
        Path to the module DOCX file
        
    Raises:
        FileNotFoundError: If module doesn't exist
    """
    # Try new storage location first
    path = STORAGE_ROOT / "system" / f"{module_name}.docx"
    if path.exists():
        return path
    
    # Fall back to legacy location
    path = MODULE_DIR / f"{module_name}.docx"
    if not path.exists():
        raise FileNotFoundError(f"Module '{module_name}' not found")
    return path


def build_context(project: Any, variables: Optional[Dict] = None) -> Dict:
    """Build the full context dict for template rendering.
    
    Args:
        project: Project object with findings
        variables: Optional user-provided variables (company_name, etc.)
        
    Returns:
        Complete context dict for Jinja2 rendering
    """
    findings_ctx = []
    
    # Risk counters
    risk_counts = {
        "Critical": 0,
        "High": 0,
        "Medium": 0,
        "Low": 0,
        "Informational": 0,
    }
    
    # SLA counters
    overdue_count = 0
    at_risk_count = 0
    on_track_count = 0
    
    for idx, f in enumerate(project.findings or [], start=1):
        risk = _normalize_risk_label(getattr(f, "risk_rating", "Low"))
        risk_counts[risk] = risk_counts.get(risk, 0) + 1
        
        # Count SLA status
        sla_status = str(getattr(f, "sla_status", "") or "")
        if "Overdue" in sla_status:
            overdue_count += 1
        elif "At Risk" in sla_status:
            at_risk_count += 1
        elif "On Track" in sla_status:
            on_track_count += 1
        
        instances = getattr(f, "instances", []) or []
        affected_resources = ", ".join([
            getattr(inst, "location", "")[:50] 
            for inst in instances[:3]
        ]) if instances else "N/A"
        if len(instances) > 3:
            affected_resources += f" (+{len(instances) - 3} more)"
        
        description_text = _strip_html(getattr(f, "description", "") or "")
        owasp_category = getattr(f, "owasp_category", None)
        status = getattr(instances[0], "status", "New - Unvalidated") if instances else "New - Unvalidated"
        
        findings_ctx.append({
            "index": idx,
            "section_number": f"1.1.{idx}",
            "risk_rating": risk,
            "title": getattr(f, "title", "Untitled"),
            "instances_count": len(instances),
            "affected_resources": affected_resources,
            "status": status,
            "owasp_vector": f"OWASP {owasp_category}" if owasp_category else "N/A",
            "description_text": description_text,  # No truncation - full description
            "remediation_text": _strip_html(getattr(f, "remediation", "") or ""),
            # Extended fields
            "impact": _strip_html(getattr(f, "impact", "") or ""),
            "references_url": getattr(f, "references_url", None) or "N/A",
            "poc_content": _strip_html(getattr(f, "poc_description", "") or ""),
            "review_status": str(getattr(f, "review_status", "Pending") or "Pending"),
            "reviewer_name": getattr(f, "reviewer_name", None) or "N/A",
            "issue_status": str(getattr(f, "issue_status", "Open") or "Open"),
            "issue_status_comment": getattr(f, "issue_status_comment", None) or "",
            "jira_issue_key": getattr(f, "jira_issue_key", None) or "N/A",
            "jira_status": getattr(f, "jira_status", None) or "N/A",
            "remediation_deadline": _fmt_dt(getattr(f, "remediation_deadline", None)),
            "sla_status": sla_status or "N/A",
            "remediation_owner": getattr(f, "remediation_owner", None) or "N/A",
            "discovered_at": _fmt_dt(getattr(f, "discovered_at", None)),
            "resolved_at": _fmt_dt(getattr(f, "resolved_at", None)),
            "owasp_category": owasp_category or "N/A",
            "cwe_id": getattr(f, "cwe_id", None) or "N/A",
            "cve_id": getattr(f, "cve_id", None) or "N/A",
            "cvss_vector": getattr(f, "cvss_vector", None) or "N/A",
            "cvss_score": getattr(f, "cvss_score", None),
            "owasp_likelihood": getattr(f, "owasp_likelihood", None),
            "owasp_impact": getattr(f, "owasp_impact", None),
            "owasp_risk_rating": getattr(f, "owasp_risk_rating", None) or "N/A",
            "template_id": getattr(f, "template_id", None) or "",
        })
    
    # Sort findings by risk (Critical -> High -> Medium -> Low -> Informational)
    RISK_ORDER = ["Critical", "High", "Medium", "Low", "Informational"]
    findings_ctx.sort(
        key=lambda x: RISK_ORDER.index(x["risk_rating"]) 
        if x["risk_rating"] in RISK_ORDER else len(RISK_ORDER)
    )
    
    # Build base context
    ctx = {
        "project": {
            "name": getattr(project, "name", "Project"),
            "consultant_name": getattr(project, "consultant_name", None) or "N/A",
        },
        "findings": findings_ctx,
        "total_findings": len(findings_ctx),
        # Risk counts
        "critical_count": risk_counts["Critical"],
        "high_count": risk_counts["High"],
        "medium_count": risk_counts["Medium"],
        "low_count": risk_counts["Low"],
        "informational_count": risk_counts["Informational"],
        # SLA counts
        "overdue_count": overdue_count,
        "at_risk_count": at_risk_count,
        "on_track_count": on_track_count,
        # Metadata
        "report_date": datetime.now().strftime("%B %d, %Y"),
        "assessment_period": "N/A",  # Can be overridden by variables
        "company_name": "N/A",  # Can be overridden by variables
    }
    
    # Merge user variables
    if variables:
        ctx.update(variables)
    
    return ctx


def render_module(module_path: Path, context: Dict, module_name: str = "") -> Document:
    """Render a single module template with the given context.
    
    Args:
        module_path: Path to the module DOCX template
        context: Full context dict for Jinja2 rendering
        module_name: Name of the module (for special handling)
        
    Returns:
        Rendered Document object
    """
    tpl = DocxTemplate(module_path)
    
    # Add donut charts for ANY template that contains findings
    # This ensures custom templates also get donut images
    if "findings" in context and context["findings"]:
        enhanced_findings = []
        for f_ctx in context["findings"]:
            risk = f_ctx.get("risk_rating", "Low")
            color = RISK_COLORS.get(risk, "DDDDDD")
            
            # Generate donut chart image
            try:
                donut_stream = _generate_donut_image(
                    risk,
                    color,
                    size_inches=1.2,  # Smaller for inline use
                    dpi=150,
                )
                donut_img = InlineImage(tpl, donut_stream, Cm(3.0))
                f_ctx["donut_img"] = donut_img
                f_ctx["has_donut"] = True
            except Exception as e:
                # Fallback to text if image generation fails
                f_ctx["donut_img"] = f"[{risk}]"
                f_ctx["has_donut"] = False
            
            enhanced_findings.append(f_ctx)
        
        # Update context with enhanced findings
        context["findings"] = enhanced_findings
    
    tpl.render(context)
    
    # Save to BytesIO and reload as Document for merging
    buf = BytesIO()
    tpl.save(buf)
    buf.seek(0)
    doc = Document(buf)
    
    # Post-process: Add colored left borders to tables in ANY template with findings
    # This ensures custom templates also get colored borders
    if "findings" in context and context["findings"]:
        # Apply colored borders to finding tables
        findings_iter = iter(context["findings"])
        for tbl in doc.tables:
            try:
                finding = next(findings_iter)
                color = RISK_COLORS.get(finding.get("risk_rating", "Low"), "DDDDDD")
                _set_table_left_border(tbl, color)
            except StopIteration:
                break
            except Exception:
                continue
    
    return doc


def merge_documents(docs: List[Document]) -> bytes:
    """Merge multiple Document objects into a single DOCX.
    
    Uses docxcompose to preserve styles, headers, footers, and images.
    
    Args:
        docs: List of Document objects to merge
        
    Returns:
        Merged DOCX as bytes
        
    Raises:
        ValueError: If docs list is empty
    """
    if not docs:
        raise ValueError("Cannot merge empty document list")
    
    if len(docs) == 1:
        # Single document - just return it
        buf = BytesIO()
        docs[0].save(buf)
        buf.seek(0)
        return buf.read()
    
    # Use docxcompose Composer to merge
    base = docs[0]
    composer = Composer(base)
    
    for doc in docs[1:]:
        composer.append(doc)
    
    # Save merged document
    out = BytesIO()
    composer.save(out)
    out.seek(0)
    return out.read()


def add_watermark_to_docx(docx_bytes: bytes, watermark_text: str = "PREVIEW - NOT FINAL") -> bytes:
    """Add watermark text to the header of a DOCX document.
    
    Args:
        docx_bytes: Original DOCX as bytes
        watermark_text: Text to display as watermark
        
    Returns:
        Modified DOCX with watermark as bytes
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Load document
    buf = BytesIO(docx_bytes)
    doc = Document(buf)
    
    # Add watermark to all sections' headers
    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            header_para = header.add_paragraph()
        else:
            header_para = header.paragraphs[0]
        
        # Clear existing content and add watermark
        header_para.clear()
        run = header_para.add_run(watermark_text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save modified document
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def assemble_report(
    session: Session,
    project: Any,
    template_ids: List[int],
    variables: Optional[Dict] = None,
) -> bytes:
    """Assemble a modular report from selected templates.
    
    Main entry point for modular report generation using database templates.
    
    Args:
        session: Database session
        project: Project object with findings
        template_ids: List of template IDs to include (in order)
        variables: Optional user variables (company_name, etc.)
        
    Returns:
        Complete assembled DOCX report as bytes
        
    Raises:
        FileNotFoundError: If a requested template doesn't exist
        ValueError: If template_ids list is empty
    """
    if not template_ids:
        raise ValueError("Must specify at least one template")
    
    # Load templates from database
    templates = []
    for tmpl_id in template_ids:
        tmpl = get_template_by_id(session, tmpl_id)
        if not tmpl:
            raise ValueError(f"Template ID {tmpl_id} not found")
        templates.append(tmpl)
    
    # Get paths and validate all templates exist
    template_paths = [get_template_path(tmpl) for tmpl in templates]
    
    # Build context once for all templates
    context = build_context(project, variables)
    
    # Render each template
    rendered_docs = []
    for tmpl, path in zip(templates, template_paths):
        try:
            # Use template name as module_name for special handling (e.g., "Detailed Findings")
            doc = render_module(path, context, module_name=tmpl.name.lower().replace(" ", "_"))
            rendered_docs.append(doc)
        except Exception as e:
            # Add context about which template failed
            raise RuntimeError(f"Failed to render template '{tmpl.name}': {e}") from e
    
    # Merge all rendered templates
    return merge_documents(rendered_docs)


def assemble_report_legacy(
    project: Any,
    modules: List[str],
    variables: Optional[Dict] = None,
) -> bytes:
    """DEPRECATED: Assemble a modular report from module names.
    
    This function is kept for backward compatibility.
    New code should use assemble_report() with template IDs.
    
    Args:
        project: Project object with findings
        modules: List of module names to include (in order)
        variables: Optional user variables (company_name, etc.)
        
    Returns:
        Complete assembled DOCX report as bytes
        
    Raises:
        FileNotFoundError: If a requested module doesn't exist
        ValueError: If modules list is empty
    """
    if not modules:
        raise ValueError("Must specify at least one module")
    
    # Validate all modules exist before rendering (use legacy path resolution)
    module_paths = [get_module_path(m) for m in modules]
    
    # Build context once for all modules
    context = build_context(project, variables)
    
    # Render each module with module name for special handling
    rendered_docs = []
    for module_name, path in zip(modules, module_paths):
        try:
            doc = render_module(path, context, module_name=module_name)
            rendered_docs.append(doc)
        except Exception as e:
            # Add context about which module failed
            raise RuntimeError(f"Failed to render module '{path.stem}': {e}") from e
    
    # Merge all rendered modules
    return merge_documents(rendered_docs)


def list_available_modules(session: Optional[Session] = None, project_id: Optional[int] = None) -> List[Dict[str, Any]]:
    """List all available report templates from database.
    
    Args:
        session: Optional database session (for API endpoint use)
        project_id: Optional project ID (for future project-specific filtering)
    
    Returns:
        List of dicts with template info (id, name, description, type, etc.)
    """
    if not session:
        # If no session provided, return empty list (caller should provide session)
        return []
    
    # Query all templates with DOCX files
    # For now, show all non-private templates (system + public + user uploads)
    # TODO: Add project_id filtering when model is updated
    stmt = select(ReportTemplate).where(
        ReportTemplate.docx_file_path != None
    )
    templates = session.exec(stmt).all()
    
    modules = []
    for tmpl in templates:
        try:
            path = get_template_path(tmpl)
            exists = path.exists()
        except (FileNotFoundError, ValueError):
            exists = False
            path = None
        
        modules.append({
            "id": tmpl.id,
            "name": tmpl.name,
            "description": tmpl.description or "",
            "template_type": str(tmpl.template_type),
            "exists": exists,
            "path": str(path) if path else None,
            "is_system": tmpl.is_system_template,
            "is_public": tmpl.is_public,
        })
    
    return modules


def _get_module_description(module_name: str) -> str:
    """Get a human-readable description for a module."""
    descriptions = {
        "title_page": "Project title, metadata, and company branding",
        "executive_summary": "High-level overview and key metrics",
        "risk_charts": "Visual risk distribution and trends",
        "top_findings": "Top N critical findings summary",
        "detailed_findings": "Full finding details with all fields",
        "recommendations": "Remediation recommendations and action items",
        "appendix": "Additional technical details and references",
        "sla_status": "SLA tracking and deadline summary",
        "compliance_owasp": "OWASP Top 10 compliance mapping",
        "compliance_cwe": "CWE Top 25 compliance mapping",
        "jira_integration": "Jira ticket status and linking",
    }
    return descriptions.get(module_name, "No description available")
